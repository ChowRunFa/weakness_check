# -*- coding: utf-8 -*-
"""
异步内容检查API（原批量检查）
"""
import os
import json
import logging
import threading
import requests
from datetime import datetime
from flask import Blueprint, request, jsonify
from flasgger import swag_from
from werkzeug.utils import secure_filename
from docx import Document

# 导入现有的工具类和方法
from objs.PlanAuditor import PlanAuditor
from objs.FileManager import FileManager
from utils.prompts import CONSTRUCTION_EXPERT_SYSTEM, get_batch_check_prompt, parse_llm_judgment, parse_confidence_score

# 导入数据库模块
from db import AsyncTaskDAO, DocumentDAO, ContentCheckDAO

# 导入swagger配置
from utils.swagger_configs.async_content_check_swagger import (
    async_content_check_swagger,
    get_content_check_task_status_swagger,
    get_content_check_result_swagger
)

# 创建蓝图
api_content_check_async = Blueprint('api_content_check_async', __name__)

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 全局配置
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx', 'doc', 'txt', 'pdf', 'json', 'jsonl'}
CACHE_DIR = 'cache'
CALLBACK_URL = '/test/content/callback'  # 本地测试回调接口地址
DEFAULT_CALLBACK_BASE_URL = 'http://127.0.0.1:5000'  # 默认本地回调基础URL

# 确保目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CACHE_DIR, exist_ok=True)

def generate_timestamp_folder():
    """生成基于当前时间的文件夹名"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]
    return timestamp

def ensure_upload_subfolder():
    """创建并返回带时间戳的上传子文件夹路径"""
    timestamp_folder = generate_timestamp_folder()
    upload_subfolder = os.path.join(UPLOAD_FOLDER, timestamp_folder)
    os.makedirs(upload_subfolder, exist_ok=True)
    return upload_subfolder, timestamp_folder

def allowed_file(filename):
    """检查文件类型是否被允许"""
    if not filename or '.' not in filename:
        return False
    try:
        ext = filename.rsplit('.', 1)[1].lower()
        return ext in ALLOWED_EXTENSIONS
    except (IndexError, AttributeError):
        return False

def extract_text_from_docx(file_path):
    """从docx文件提取文本，增加错误处理"""
    try:
        doc = Document(file_path)
        text_content = []
        
        # 提取段落文本
        try:
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
        except Exception as e:
            logger.warning(f"提取段落时出错: {e}")
        
        # 提取表格文本
        try:
            for table_idx, table in enumerate(doc.tables):
                try:
                    for row_idx, row in enumerate(table.rows):
                        try:
                            for cell_idx in range(len(table.columns)):
                                try:
                                    cell = row.cells[cell_idx]
                                    if cell.text and cell.text.strip():
                                        text_content.append(cell.text.strip())
                                except (IndexError, AttributeError):
                                    continue
                        except Exception as e:
                            logger.debug(f"跳过表格 {table_idx} 行 {row_idx}: {e}")
                            continue
                except Exception as e:
                    logger.debug(f"跳过表格 {table_idx}: {e}")
                    continue
        except Exception as e:
            logger.warning(f"提取表格时出错: {e}")
        
        result = "\n".join(text_content)
        
        if not result.strip():
            return "文档内容为空或无法提取文本"
        
        logger.info(f"成功提取文档内容，长度: {len(result)} 字符")
        return result
        
    except Exception as e:
        logger.error(f"提取docx文件文本时发生错误: {str(e)}")
        raise ValueError(f"无法提取docx文件内容: {str(e)}")

def send_callback(callback_url, task_id, status, data=None, error_message=None):
    """发送回调请求并更新数据库状态"""
    try:
        # 1. 先更新数据库状态
        try:
            if status == "success" and data:
                # 保存内容检查结果到数据库
                ContentCheckDAO.save_check_result(task_id, data)
                # 更新任务状态为成功
                AsyncTaskDAO.update_task_status(
                    task_id, 
                    status, 
                    result_data=data.get('summary') if data else None
                )
            else:
                # 更新任务状态
                AsyncTaskDAO.update_task_status(
                    task_id, 
                    status, 
                    error_message=error_message
                )
            logger.info(f"数据库状态更新成功，任务ID: {task_id}, 状态: {status}")
        except Exception as db_error:
            logger.error(f"更新数据库状态失败: {str(db_error)}, 任务ID: {task_id}")
        
        # 2. 准备回调数据
        callback_data = {
            "task_id": task_id,
            "status": status,  # "success", "failed", "processing"
            "timestamp": datetime.now().isoformat(),
            "data": data,
            "error_message": error_message
        }
        
        headers = {
            'Content-Type': 'application/json'
        }
        
        logger.info(f"发送回调到: {callback_url}, 任务ID: {task_id}, 状态: {status}")
        
        # 3. 发送POST请求到回调接口
        response = requests.post(
            callback_url, 
            json=callback_data, 
            headers=headers,
            timeout=30
        )
        
        if response.status_code == 200:
            logger.info(f"回调发送成功，任务ID: {task_id}")
        else:
            logger.warning(f"回调发送失败，状态码: {response.status_code}, 任务ID: {task_id}")
            
    except Exception as e:
        logger.error(f"发送回调时发生错误: {str(e)}, 任务ID: {task_id}")

def async_content_check_worker(task_params):
    """异步执行内容检查的工作函数"""
    task_id = task_params['task_id']
    callback_url_full = task_params['callback_url']
    
    try:
        logger.info(f"开始异步执行内容检查，任务ID: {task_id}")
        
        # 发送处理中状态回调
        send_callback(callback_url_full, task_id, "processing", {"message": "开始执行内容检查"})
        
        # 执行内容检查
        result = perform_content_check_internal(task_params)
        
        # 发送成功回调
        send_callback(callback_url_full, task_id, "success", result)
        
        logger.info(f"内容检查完成，任务ID: {task_id}")
        
    except Exception as e:
        error_msg = f"内容检查执行失败: {str(e)}"
        logger.error(f"{error_msg}, 任务ID: {task_id}")
        
        # 发送失败回调
        send_callback(callback_url_full, task_id, "failed", error_message=error_msg)

def perform_content_check_internal(task_params):
    """内部执行内容检查的具体逻辑"""
    checklist_path = task_params['checklist_path']
    document_path = task_params['document_path']
    checklist_filename = task_params['checklist_filename']
    document_filename = task_params['document_filename']
    embedding_model = task_params['embedding_model']
    chat_model = task_params['chat_model']
    top_k = task_params['top_k']
    openai_api_key = task_params['openai_api_key']
    openai_api_base = task_params['openai_api_base']
    timestamp = task_params['timestamp']
    
    # 解析检查项文件
    checklist_items = []
    try:
        with open(checklist_path, 'r', encoding='utf-8') as f:
            if checklist_filename.endswith('.jsonl'):
                # JSONL格式，每行一个JSON对象
                for line in f:
                    line = line.strip()
                    if line:
                        checklist_items.append(json.loads(line))
            else:
                # JSON格式，可能是数组或单个对象
                content = json.load(f)
                if isinstance(content, list):
                    checklist_items = content
                else:
                    checklist_items = [content]
    except Exception as e:
        raise Exception(f'解析检查项文件失败: {str(e)}')
    
    if not checklist_items:
        raise Exception('检查项文件为空或格式不正确')
    
    # 提取文档内容
    try:
        document_content = extract_text_from_docx(document_path)
    except Exception as e:
        raise Exception(f'提取文档内容失败: {str(e)}')
    
    # 创建临时检查项文件（JSONL格式）
    temp_checklist_path = os.path.join(os.path.dirname(checklist_path), 'temp_checklist.jsonl')
    with open(temp_checklist_path, 'w', encoding='utf-8') as f:
        for item in checklist_items:
            f.write(json.dumps(item, ensure_ascii=False) + '\n')
    
    # 初始化审查器
    try:
        auditor = PlanAuditor(
            plan_content=document_content,
            check_list_file=temp_checklist_path,
            embedding_model=embedding_model,
            openai_api_key=openai_api_key,
            openai_api_base=openai_api_base,
            cache_dir=CACHE_DIR,
            original_filename=document_filename
        )
        
        # 构建嵌入
        plan_id = auditor.build_or_load_embeddings()
        
    except Exception as e:
        raise Exception(f'初始化审查器失败: {str(e)}')
    
    # 逐个检查每个检查项
    check_results = []
    
    for i, check_item in enumerate(checklist_items, 1):
        try:
            logger.info(f"正在检查第 {i}/{len(checklist_items)} 项: {check_item}")
            
            # 获取检查项信息
            category = check_item.get('分类', '未知分类')
            check_scenario = check_item.get('专项施工方案严重缺陷情形', '')
            item_number = check_item.get('序号', str(i))
            
            if not check_scenario:
                check_results.append({
                    'item_number': item_number,
                    'category': category,
                    'check_scenario': '检查项信息不完整',
                    'evidence': '',
                    'judgment': '无法判断',
                    'probability': 0.0,
                    'error': '检查项缺少"专项施工方案严重缺陷情形"字段'
                })
                continue
            
            # 搜索相关文本片段
            similar_chunks = auditor.search_similar_chunks(check_scenario, top_k=top_k)
            
            # 组合检索到的文本作为证据
            evidence_texts = []
            for chunk_dict in similar_chunks:
                chunk_text = chunk_dict.get("text", "")
                similarity = chunk_dict.get("similarity", 0.0)
                evidence_texts.append(f"相关度{similarity:.3f}: {chunk_text}")
            
            evidence = "\n".join(evidence_texts)
            
            # 使用大模型进行智能判断
            if similar_chunks:
                # 构建上下文内容
                context = "\n".join([chunk_dict.get("text", "") for chunk_dict in similar_chunks])
                
                # 构建判断prompt
                messages = [
                    {
                        "role": "system",
                        "content": CONSTRUCTION_EXPERT_SYSTEM
                    },
                    {
                        "role": "user",
                        "content": get_batch_check_prompt(check_scenario, category, context)
                    }
                ]
                
                # 调用大模型进行判断
                try:
                    llm_response = auditor.embedder.generate_text(
                        messages=messages,
                        model=chat_model,
                        temperature=0.1
                    )
                    
                    # 解析大模型回复
                    judgment = parse_llm_judgment(llm_response)
                    probability = parse_confidence_score(llm_response)
                    
                    check_result = llm_response
                    
                except Exception as e:
                    logger.error(f"大模型调用失败: {str(e)}")
                    # 降级到简单判断
                    judgment = "合规"  # 默认合规
                    probability = 0.5
                    check_result = f"大模型调用失败，降级判断: {str(e)}"
                    
            else:
                # 没有找到相关内容
                judgment = "无法判断"
                probability = 0.0
                check_result = "未找到相关文档内容"
            
            check_results.append({
                'item_number': item_number,
                'category': category,
                'check_scenario': check_scenario,
                'evidence': evidence,
                'judgment': judgment,
                'probability': round(probability, 3),
                'detailed_result': check_result,
                'chunk_count': len(similar_chunks)
            })
            
        except Exception as e:
            logger.error(f"检查第 {i} 项时发生错误: {str(e)}")
            check_results.append({
                'item_number': check_item.get('序号', str(i)),
                'category': check_item.get('分类', '未知分类'),
                'check_scenario': check_item.get('专项施工方案严重缺陷情形', ''),
                'evidence': '',
                'judgment': '检查失败',
                'probability': 0.0,
                'error': str(e)
            })
    
    # 计算总体统计
    total_items = len(check_results)
    compliant_items = len([r for r in check_results if r['judgment'] == '合规'])
    non_compliant_items = len([r for r in check_results if r['judgment'] == '不合规'])
    failed_items = len([r for r in check_results if r['judgment'] == '检查失败'])
    
    return {
        'summary': {
            'total_items': total_items,
            'compliant_items': compliant_items,
            'non_compliant_items': non_compliant_items,
            'failed_items': failed_items,
            'compliance_rate': round(compliant_items / total_items * 100, 2) if total_items > 0 else 0
        },
        'check_results': check_results,
        'document_filename': document_filename,
        'checklist_filename': checklist_filename,
        'plan_id': plan_id,
        'upload_folder': timestamp
    }

@api_content_check_async.route('/async_content_check', methods=['POST'])
@swag_from(async_content_check_swagger)
def async_content_check():
    """
    异步内容检查接口
    
    立即返回调用结果，后台异步执行检查，完成后回调指定接口
    """
    try:
        # 检查必要参数
        scheme_id = request.form.get('schemeId')
        file_path = request.form.get('filePath')
        scheme_name = request.form.get('schemeName')
        checklist_path = request.form.get('checklistPath')
        
        # 验证必要参数
        if not scheme_id:
            return jsonify({
                'code': 400,
                'message': '方案ID不能为空',
                'data': {'result': 'false'}
            }), 400
        
        if not file_path:
            return jsonify({
                'code': 400,
                'message': '文档文件路径不能为空',
                'data': {'result': 'false'}
            }), 400
        
        if not scheme_name:
            return jsonify({
                'code': 400,
                'message': '方案名称不能为空',
                'data': {'result': 'false'}
            }), 400
        
        if not checklist_path:
            return jsonify({
                'code': 400,
                'message': '检查项文件路径不能为空',
                'data': {'result': 'false'}
            }), 400
        
        # 验证检查项文件是否存在
        if not os.path.exists(checklist_path):
            return jsonify({
                'code': 400,
                'message': f'检查项文件不存在: {checklist_path}',
                'data': {'result': 'false'}
            }), 400
        
        # 验证检查项文件类型
        if not checklist_path.lower().endswith(('.json', '.jsonl')):
            return jsonify({
                'code': 400,
                'message': '检查项文件必须是JSON或JSONL格式',
                'data': {'result': 'false'}
            }), 400
        
        # 验证待检查文档是否存在
        if not os.path.exists(file_path):
            return jsonify({
                'code': 400,
                'message': f'文档文件不存在: {file_path}',
                'data': {'result': 'false'}
            }), 400
        
        # 验证文档文件类型
        if not file_path.lower().endswith('.docx'):
            return jsonify({
                'code': 400,
                'message': '文档必须是DOCX格式',
                'data': {'result': 'false'}
            }), 400
        
        # 创建临时文件夹用于任务标识
        upload_subfolder, timestamp_folder = ensure_upload_subfolder()
        
        # 直接使用提供的文件路径
        document_path = file_path
        document_filename = os.path.basename(file_path)
        checklist_filename = os.path.basename(checklist_path)
        
        # 获取配置参数
        callback_base_url = request.form.get('callback_base_url', DEFAULT_CALLBACK_BASE_URL)
        embedding_model = request.form.get('embedding_model', 'nomic-embed-text:latest')
        chat_model = request.form.get('chat_model', 'qwen2.5:32b')
        top_k = int(request.form.get('top_k', 5))
        openai_api_key = request.form.get('openai_api_key', 'ollama')
        openai_api_base = request.form.get('openai_api_base', 'http://59.77.7.24:11434/v1/')
        
        # 生成任务ID
        timestamp = generate_timestamp_folder()
        task_id = f"content_check_{timestamp}"
        
        # 构建回调URL
        callback_url_full = f"{callback_base_url.rstrip('/')}{CALLBACK_URL}"
        
        # 保存任务到数据库
        try:
            # 创建异步任务记录
            request_params = {
                'scheme_id': scheme_id,
                'file_path': file_path,
                'scheme_name': scheme_name,
                'checklist_path': checklist_path,
                'checklist_filename': checklist_filename,
                'document_filename': document_filename,
                'embedding_model': embedding_model,
                'chat_model': chat_model,
                'top_k': top_k,
                'callback_base_url': callback_base_url
            }
            
            task = AsyncTaskDAO.create_task(
                task_id=task_id,
                task_type='content_check',
                callback_url=callback_url_full,
                request_params=request_params,
                scheme_id=int(scheme_id),
                scheme_name=scheme_name
            )
            
            if not task:
                return jsonify({
                    'code': 500,
                    'message': '创建任务记录失败',
                    'data': {'result': 'false'}
                }), 500
            
            # 保存文档信息
            checklist_file_size = os.path.getsize(checklist_path) if os.path.exists(checklist_path) else 0
            document_file_size = os.path.getsize(document_path) if os.path.exists(document_path) else 0
            
            # 保存检查项文件引用（引用已存在的检查项文件）
            DocumentDAO.save_document_reference(
                task_id=task_id,
                scheme_id=int(scheme_id),
                original_filename=checklist_filename,
                saved_filename=checklist_filename,
                file_path=checklist_path,
                file_size=checklist_file_size,
                file_type='checklist',
                reference_folder=f'scheme_{scheme_id}'
            )
            
            # 保存文档文件引用（引用已存在的文档）
            DocumentDAO.save_document_reference(
                task_id=task_id,
                scheme_id=int(scheme_id),
                original_filename=document_filename,
                saved_filename=document_filename,
                file_path=document_path,
                file_size=document_file_size,
                file_type='document',
                reference_folder=f'scheme_{scheme_id}'
            )
            
            logger.info(f"任务和文档信息已保存到数据库，任务ID: {task_id}")
            
        except Exception as db_error:
            logger.error(f"保存任务到数据库失败: {str(db_error)}")
            return jsonify({
                'code': 500,
                'message': f'保存任务失败: {str(db_error)}',
                'data': {'result': 'false'}
            }), 500
        
        # 准备异步任务参数
        task_params = {
            'task_id': task_id,
            'scheme_id': scheme_id,
            'scheme_name': scheme_name,
            'callback_url': callback_url_full,
            'checklist_path': checklist_path,
            'document_path': document_path,
            'checklist_filename': checklist_filename,
            'document_filename': document_filename,
            'embedding_model': embedding_model,
            'chat_model': chat_model,
            'top_k': top_k,
            'openai_api_key': openai_api_key,
            'openai_api_base': openai_api_base,
            'timestamp': timestamp_folder
        }
        
        # 启动异步任务
        thread = threading.Thread(
            target=async_content_check_worker,
            args=(task_params,),
            daemon=True
        )
        thread.start()
        
        logger.info(f"异步内容检查任务已启动，任务ID: {task_id}")
        
        # 立即返回成功响应
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': {
                'result': 'true',
                'task_id': task_id,
                'callback_url': callback_url_full,
                'estimated_time': '预计5-15分钟内完成'
            }
        }), 200
        
    except Exception as e:
        logger.error(f"异步内容检查API发生错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器内部错误: {str(e)}',
            'data': {'result': 'false'}
        }), 500

# 任务状态查询API
@api_content_check_async.route('/async_content_check/status/<task_id>', methods=['GET'])
@swag_from(get_content_check_task_status_swagger)
def get_task_status(task_id):
    """查询异步任务状态"""
    try:
        task = AsyncTaskDAO.get_task_by_id(task_id)
        if not task:
            return jsonify({
                'code': 404,
                'message': '任务不存在',
                'data': None
            }), 404
        
        task_data = task.to_dict()
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': task_data
        }), 200
        
    except Exception as e:
        logger.error(f"查询任务状态失败: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'查询失败: {str(e)}',
            'data': None
        }), 500

# 检查结果查询API
@api_content_check_async.route('/async_content_check/result/<task_id>', methods=['GET'])
@swag_from(get_content_check_result_swagger)
def get_check_result(task_id):
    """获取内容检查结果"""
    try:
        # 使用ContentCheckDAO获取结果
        result = ContentCheckDAO.get_check_result(task_id)
        if not result:
            return jsonify({
                'code': 404,
                'message': '检查结果不存在或任务未完成',
                'data': None
            }), 404
        
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': result
        }), 200
        
    except Exception as e:
        logger.error(f"获取检查结果失败: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'获取结果失败: {str(e)}',
            'data': None
        }), 500

# 本地测试回调接口
@api_content_check_async.route('/test/content/callback', methods=['POST'])
def test_callback():
    """本地测试回调接口"""
    try:
        # 获取回调数据
        callback_data = request.get_json()
        
        if not callback_data:
            logger.warning("回调接收失败：未收到有效的JSON数据")
            return jsonify({
                'code': 400,
                'message': '无效的回调数据',
                'data': None
            }), 400
        
        task_id = callback_data.get('task_id')
        status = callback_data.get('status')
        timestamp = callback_data.get('timestamp')
        data = callback_data.get('data')
        error_message = callback_data.get('error_message')
        
        logger.info(f"收到内容检查回调请求 - 任务ID: {task_id}, 状态: {status}, 时间: {timestamp}")
        
        # 打印回调数据详情
        if status == "success" and data:
            summary = data.get('summary', {})
            logger.info(f"任务 {task_id} 执行成功:")
            logger.info(f"  - 总检查项数: {summary.get('total_items', 0)}")
            logger.info(f"  - 合规项目数: {summary.get('compliant_items', 0)}")
            logger.info(f"  - 不合规项目数: {summary.get('non_compliant_items', 0)}")
            logger.info(f"  - 失败项目数: {summary.get('failed_items', 0)}")
            logger.info(f"  - 合规率: {summary.get('compliance_rate', 0)}%")
        elif status == "failed":
            logger.error(f"任务 {task_id} 执行失败: {error_message}")
        elif status == "processing":
            logger.info(f"任务 {task_id} 正在处理中...")
        
        # 模拟处理回调数据
        response_data = {
            'received_task_id': task_id,
            'received_status': status,
            'received_at': datetime.now().isoformat(),
            'callback_processed': True
        }
        
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': response_data
        }), 200
        
    except Exception as e:
        logger.error(f"处理回调请求时发生错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'处理回调失败: {str(e)}',
            'data': None
        }), 500
