# -*- coding: utf-8 -*-
"""
施工方案审核API
"""
import os
import json
import logging
from datetime import datetime, timedelta
from flask import Blueprint, request, jsonify, Response
from flasgger import swag_from
from werkzeug.utils import secure_filename
from docx import Document
from objs.PlanAuditor import PlanAuditor
from objs.FileManager import FileManager
from objs.EmbeddingRetriever import EmbeddingRetriever
from utils.prompts import (
    CONSTRUCTION_EXPERT_SYSTEM, 
    get_batch_check_prompt,
    get_category_check_prompt,
    get_citation_check_prompt,
    get_structure_check_prompt,
    get_query_prompt,
    parse_llm_judgment,
    parse_confidence_score,
    generate_single_check_prompt,
    generate_category_prompt,
    generate_query_prompt,
    generate_retrieval_prompt
)
from utils.swagger_configs.ra_check_swagger import (
    upload_plan_swagger,
    query_plan_swagger, 
    check_category_swagger,
    full_audit_swagger,
    get_status_swagger,
    simple_search_swagger,
    stream_query_swagger,
    available_embeddings_swagger,
    list_files_swagger,
    delete_file_swagger,
    list_upload_folders_swagger,
    cleanup_uploads_swagger,
    batch_check_swagger,
    cite_check_swagger,
    structure_check_swagger
)

# 创建蓝图
api_ra_check = Blueprint('api_ra_check', __name__)

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 全局配置
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx', 'doc', 'txt', 'pdf'}
CACHE_DIR = 'cache'

# 确保目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CACHE_DIR, exist_ok=True)

def generate_timestamp_folder():
    """生成基于当前时间的文件夹名"""
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]  # 精确到毫秒
    return timestamp

def ensure_upload_subfolder():
    """创建并返回带时间戳的上传子文件夹路径"""
    timestamp_folder = generate_timestamp_folder()
    upload_subfolder = os.path.join(UPLOAD_FOLDER, timestamp_folder)
    os.makedirs(upload_subfolder, exist_ok=True)
    return upload_subfolder, timestamp_folder

def cleanup_old_upload_folders(days_old=7):
    """清理超过指定天数的上传文件夹"""
    if not os.path.exists(UPLOAD_FOLDER):
        return 0
    
    from datetime import datetime
    cleaned_count = 0
    cutoff_time = datetime.now().timestamp() - (days_old * 24 * 60 * 60)
    
    try:
        for item in os.listdir(UPLOAD_FOLDER):
            item_path = os.path.join(UPLOAD_FOLDER, item)
            if os.path.isdir(item_path):
                # 检查文件夹的创建时间
                folder_ctime = os.path.getctime(item_path)
                if folder_ctime < cutoff_time:
                    try:
                        import shutil
                        shutil.rmtree(item_path)
                        logger.info(f"清理旧的上传文件夹: {item}")
                        cleaned_count += 1
                    except Exception as e:
                        logger.warning(f"清理文件夹 {item} 失败: {e}")
    except Exception as e:
        logger.error(f"清理上传文件夹时发生错误: {e}")
    
    return cleaned_count

def allowed_file(filename):
    """检查文件类型是否被允许"""
    if not filename or '.' not in filename:
        return False
    try:
        ext = filename.rsplit('.', 1)[1].lower()
        return ext in ALLOWED_EXTENSIONS
    except (IndexError, AttributeError):
        return False

def extract_text_from_docx(file_path):
    """从docx文件提取文本，增加错误处理"""
    try:
        doc = Document(file_path)
        text_content = []
        
        # 提取段落文本
        try:
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
        except Exception as e:
            logger.warning(f"提取段落时出错: {e}")
        
        # 提取表格文本（使用更安全的方法）
        try:
            for table_idx, table in enumerate(doc.tables):
                try:
                    # 使用更安全的方式访问表格
                    for row_idx, row in enumerate(table.rows):
                        try:
                            # 直接访问cell内容，避免使用row.cells属性
                            for cell_idx in range(len(table.columns)):
                                try:
                                    cell = row.cells[cell_idx]
                                    if cell.text and cell.text.strip():
                                        text_content.append(cell.text.strip())
                                except (IndexError, AttributeError) as e:
                                    # 跳过有问题的单元格
                                    continue
                        except Exception as e:
                            # 跳过有问题的行
                            logger.debug(f"跳过表格 {table_idx} 行 {row_idx}: {e}")
                            continue
                except Exception as e:
                    # 跳过有问题的表格
                    logger.debug(f"跳过表格 {table_idx}: {e}")
                    continue
        except Exception as e:
            logger.warning(f"提取表格时出错: {e}")
        
        result = "\n".join(text_content)
        
        # 检查是否提取到内容
        if not result.strip():
            return "文档内容为空或无法提取文本"
        
        logger.info(f"成功提取文档内容，长度: {len(result)} 字符")
        return result
        
    except Exception as e:
        logger.error(f"提取docx文件文本时发生错误: {str(e)}")
        # 如果完全失败，尝试只提取段落
        try:
            doc = Document(file_path)
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            if text_content:
                result = "\n".join(text_content)
                logger.info(f"降级提取成功，仅段落内容，长度: {len(result)} 字符")
                return result
            else:
                return "文档仅包含表格内容，无法安全提取"
        except Exception as fallback_e:
            logger.error(f"降级提取也失败: {fallback_e}")
            raise ValueError(f"无法提取docx文件内容: {str(e)}")

@api_ra_check.route('/ra_check/upload_plan', methods=['POST'])
@swag_from(upload_plan_swagger)
def upload_plan():
    """上传施工方案文档"""
    try:
        if 'file' not in request.files:
            return jsonify({'status': 'error', 'message': '未找到文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'status': 'error', 'message': '未选择文件'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'status': 'error', 'message': '不支持的文件类型'}), 400
        
        # 创建时间戳子文件夹并保存文件
        upload_subfolder, timestamp_folder = ensure_upload_subfolder()
        filename = secure_filename(file.filename)
        file_path = os.path.join(upload_subfolder, filename)
        file.save(file_path)
        
        # 记录文件的相对路径信息
        relative_file_path = os.path.join(timestamp_folder, filename)
        
        # 提取文本内容
        if filename.endswith('.docx'):
            plan_content = extract_text_from_docx(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                plan_content = f.read()
        
        # 获取配置参数
        embedding_model = request.form.get('embedding_model', 'nomic-embed-text:latest')
        openai_api_key = request.form.get('openai_api_key', 'ollama')
        openai_api_base = request.form.get('openai_api_base', 'http://59.77.7.24:11434/v1/')
        
        # 验证文本内容
        if not plan_content or not plan_content.strip():
            return jsonify({'status': 'error', 'message': '文档内容为空'}), 400
        
        # 创建审核器，传入原始文件名
        auditor = PlanAuditor(
            plan_content=plan_content,
            check_list_file='data/checklist/weakness_list.jsonl',
            embedding_model=embedding_model,
            openai_api_key=openai_api_key,
            openai_api_base=openai_api_base,
            cache_dir=CACHE_DIR,
            original_filename=file.filename  # 传入原始文件名
        )
        
        # 构建嵌入
        plan_id = auditor.build_or_load_embeddings()
        
        # 保存auditor实例到缓存
        if not hasattr(api_ra_check, 'auditor_cache'):
            api_ra_check.auditor_cache = {}
        api_ra_check.auditor_cache[plan_id] = auditor
        
        return jsonify({
            'status': 'success',
            'message': '文档上传成功',
            'plan_id': plan_id,
            'original_filename': file.filename,
            'saved_path': relative_file_path,
            'timestamp_folder': timestamp_folder,
            'text_length': len(plan_content),
            'chunks_count': len(auditor.chunks),
            'embedding_model': embedding_model
        }), 200
        
    except Exception as e:
        logger.error(f"上传文档时发生错误: {str(e)}")
        return jsonify({'status': 'error', 'message': f'处理文档时发生错误: {str(e)}'}), 500

@api_ra_check.route('/ra_check/query', methods=['POST'])
@swag_from(query_plan_swagger)
def query_plan():
    """查询方案内容"""
    try:
        data = request.get_json()
        plan_id = data.get('plan_id')
        query = data.get('query')
        top_k = data.get('top_k', 5)
        stream = data.get('stream', False)
        
        if not plan_id or not query:
            return jsonify({'status': 'error', 'message': '缺少必要参数'}), 400
        
        # 从缓存中获取auditor
        if not hasattr(api_ra_check, 'auditor_cache') or plan_id not in api_ra_check.auditor_cache:
            return jsonify({'status': 'error', 'message': '方案未找到，请先上传方案'}), 404
        
        auditor = api_ra_check.auditor_cache[plan_id]
        
        if stream:
            # 流式输出
            def generate_stream_response():
                try:
                    # 发送开始信号
                    yield f"data: {json.dumps({'type': 'start', 'message': '开始查询...'})}\n\n"
                    
                    # 搜索相似文本块
                    similar_chunks = auditor.search_similar_chunks(query, top_k)
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'找到 {len(similar_chunks)} 个相关文本块，正在生成回复...'})}\n\n"
                    
                    # 构建检索上下文
                    context_texts = []
                    for chunk_result in similar_chunks:
                        context_texts.append(chunk_result.get('text', ''))
                    context = "\n\n".join(context_texts)
                    
                    # 构建prompt
                    messages = [
                        {"role": "system", "content": "你是一个专业的施工方案审核助手。请基于提供的施工方案内容，回答用户的问题。请严格基于提供的施工方案内容进行回答，如果方案中没有相关信息，请明确说明。"},
                        {"role": "user", "content": f"施工方案内容：\n{context}\n\n用户问题：{query}\n\n请基于上述施工方案内容回答用户问题："}
                    ]
                    
                    # 流式生成回复
                    yield f"data: {json.dumps({'type': 'generation_start', 'message': '正在生成回复...'})}\n\n"
                    
                    full_response = ""
                    for token in auditor.embedder.generate_text_stream(messages, model='qwen2.5:7b'):
                        full_response += token
                        yield f"data: {json.dumps({'type': 'token', 'content': token})}\n\n"
                    
                    # 发送完成信号
                    yield f"data: {json.dumps({'type': 'complete', 'message': '生成完成', 'full_response': full_response, 'context_chunks_count': len(similar_chunks)})}\n\n"
                    
                except Exception as e:
                    logger.error(f"流式查询发生错误: {str(e)}")
                    yield f"data: {json.dumps({'type': 'error', 'message': f'查询发生错误: {str(e)}'})}\n\n"
            
            return Response(
                generate_stream_response(),
                mimetype='text/plain',
                headers={
                    'Cache-Control': 'no-cache',
                    'Connection': 'keep-alive',
                    'Access-Control-Allow-Origin': '*',
                    'Access-Control-Allow-Headers': 'Content-Type'
                }
            )
        else:
            # 非流式输出
            results = auditor.response_user_query(query, top_k)
            return jsonify({
                'status': 'success',
                'results': results
            }), 200
        
    except Exception as e:
        logger.error(f"查询时发生错误: {str(e)}")
        return jsonify({'status': 'error', 'message': f'查询时发生错误: {str(e)}'}), 500

@api_ra_check.route('/ra_check/check_category', methods=['POST'])
@swag_from(check_category_swagger)
def check_category():
    """分类场景检查"""
    try:
        data = request.get_json()
        plan_id = data.get('plan_id')
        category = data.get('category')
        scenario = data.get('scenario')
        top_k = data.get('top_k', 5)
        stream = data.get('stream', False)
        
        if not plan_id or not category or not scenario:
            return jsonify({'status': 'error', 'message': '缺少必要参数'}), 400
        
        # 从缓存中获取auditor
        if not hasattr(api_ra_check, 'auditor_cache') or plan_id not in api_ra_check.auditor_cache:
            return jsonify({'status': 'error', 'message': '方案未找到，请先上传方案'}), 404
        
        auditor = api_ra_check.auditor_cache[plan_id]
        
        if stream:
            # 流式输出
            def generate_stream_response():
                try:
                    yield f"data: {json.dumps({'type': 'start', 'message': f'开始{category}分类检查...'})}\n\n"
                    
                    # 搜索相关文本块
                    similar_chunks = auditor.search_similar_chunks(scenario, top_k)
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'找到 {len(similar_chunks)} 个相关文本块，正在分析...'})}\n\n"
                    
                    # 构建检索上下文
                    context_texts = []
                    for chunk_result in similar_chunks:
                        context_texts.append(chunk_result.get('text', ''))
                    context = "\n\n".join(context_texts)
                    
                    # 构建prompt
                    messages = [
                        {"role": "system", "content": f"你是一个专业的{category}施工方案审核专家。请基于提供的施工方案内容，针对特定缺陷情形进行专业分析。"},
                        {"role": "user", "content": f"施工方案内容：\n{context}\n\n需要检查的缺陷情形：{scenario}\n\n请基于上述内容分析是否存在该缺陷情形，并给出专业建议："}
                    ]
                    
                    # 流式生成分析结果
                    yield f"data: {json.dumps({'type': 'generation_start', 'message': '正在生成分析结果...'})}\n\n"
                    
                    full_response = ""
                    for token in auditor.embedder.generate_text_stream(messages, model='qwen2.5:7b'):
                        full_response += token
                        yield f"data: {json.dumps({'type': 'token', 'content': token})}\n\n"
                    
                    # 组装检查结果
                    plan_content = [chunk.get('text', '') for chunk in similar_chunks]
                    check_items = [{'similarity': chunk.get('similarity', 0), 'text': chunk.get('text', '')} for chunk in similar_chunks]
                    
                    # 发送完成信号
                    yield f"data: {json.dumps({'type': 'complete', 'message': '分析完成', 'analysis_result': full_response, 'plan_content': plan_content, 'check_items': check_items})}\n\n"
                    
                except Exception as e:
                    logger.error(f"流式分类检查发生错误: {str(e)}")
                    yield f"data: {json.dumps({'type': 'error', 'message': f'检查发生错误: {str(e)}'})}\n\n"
            
            return Response(
                generate_stream_response(),
                mimetype='text/plain',
                headers={
                    'Cache-Control': 'no-cache',
                    'Connection': 'keep-alive',
                    'Access-Control-Allow-Origin': '*',
                    'Access-Control-Allow-Headers': 'Content-Type'
                }
            )
        else:
            # 非流式输出
            results = auditor.check_category_scenario(category, scenario, top_k)
            return jsonify({
                'status': 'success',
                'plan_content': results['plan_content'],
                'check_items': results['check_items']
            }), 200
        
    except Exception as e:
        logger.error(f"分类检查时发生错误: {str(e)}")
        return jsonify({'status': 'error', 'message': f'分类检查时发生错误: {str(e)}'}), 500

@api_ra_check.route('/ra_check/full_audit', methods=['POST'])
@swag_from(full_audit_swagger)
def full_audit():
    """完整审核"""
    try:
        data = request.get_json()
        plan_id = data.get('plan_id')
        check_categories = data.get('check_categories', [])
        stream = data.get('stream', False)
        
        if not plan_id:
            return jsonify({'status': 'error', 'message': '缺少方案ID'}), 400
        
        # 从缓存中获取auditor
        if not hasattr(api_ra_check, 'auditor_cache') or plan_id not in api_ra_check.auditor_cache:
            return jsonify({'status': 'error', 'message': '方案未找到，请先上传方案'}), 404
        
        auditor = api_ra_check.auditor_cache[plan_id]
        
        # 获取检查项
        check_items = auditor.check_items
        
        if check_categories:
            # 过滤指定类别的检查项
            filtered_items = [item for item in check_items if item.get('分类') in check_categories]
        else:
            filtered_items = check_items
        
        if stream:
            # 流式输出
            def generate_stream_response():
                try:
                    yield f"data: {json.dumps({'type': 'start', 'message': f'开始完整审核，共{len(filtered_items)}个检查项...'})}\n\n"
                    
                    audit_results = {}
                    compliant_count = 0
                    non_compliant_count = 0
                    
                    # 对每个检查项进行审核
                    for i, item in enumerate(filtered_items, 1):
                        category = item.get('分类', '未知类别')
                        scenario = item.get('专项施工方案严重缺陷情形', '')
                        
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'正在检查第{i}/{len(filtered_items)}项: {category} - {scenario}'})}\n\n"
                        
                        # 搜索相关文本块
                        similar_chunks = auditor.search_similar_chunks(scenario, top_k=3)
                        
                        # 构建分析prompt
                        context_texts = []
                        for chunk_result in similar_chunks:
                            context_texts.append(chunk_result.get('text', ''))
                        context = "\n\n".join(context_texts)
                        
                        messages = [
                            {"role": "system", "content": f"你是一个专业的{category}施工方案审核专家。请判断方案是否合规。"},
                            {"role": "user", "content": f"施工方案内容：\n{context}\n\n缺陷情形：{scenario}\n\n请判断是否存在该缺陷并给出合规性分析："}
                        ]
                        
                        # 流式生成分析
                        analysis_result = ""
                        for token in auditor.embedder.generate_text_stream(messages, model='qwen2.5:7b'):
                            analysis_result += token
                            yield f"data: {json.dumps({'type': 'token', 'content': token, 'item_index': i, 'category': category})}\n\n"
                        
                        # 简单的合规性判断
                        is_compliant = len(similar_chunks) > 0 and "合规" in analysis_result
                        
                        if is_compliant:
                            compliant_count += 1
                        else:
                            non_compliant_count += 1
                        
                        audit_results[f"{category}_{item.get('序号', '')}"] = {
                            'category': category,
                            'scenario': scenario,
                            'plan_content': [chunk.get('text', '') for chunk in similar_chunks],
                            'check_items': similar_chunks,
                            'is_compliant': is_compliant,
                            'analysis_result': analysis_result
                        }
                        
                        yield f"data: {json.dumps({'type': 'item_complete', 'item_index': i, 'category': category, 'is_compliant': is_compliant})}\n\n"
                    
                    # 发送完成信号
                    summary = {
                        'total_checks': len(filtered_items),
                        'compliant_count': compliant_count,
                        'non_compliant_count': non_compliant_count
                    }
                    
                    yield f"data: {json.dumps({'type': 'complete', 'message': '完整审核完成', 'audit_results': audit_results, 'summary': summary})}\n\n"
                    
                except Exception as e:
                    logger.error(f"流式完整审核发生错误: {str(e)}")
                    yield f"data: {json.dumps({'type': 'error', 'message': f'审核发生错误: {str(e)}'})}\n\n"
            
            return Response(
                generate_stream_response(),
                mimetype='text/plain',
                headers={
                    'Cache-Control': 'no-cache',
                    'Connection': 'keep-alive',
                    'Access-Control-Allow-Origin': '*',
                    'Access-Control-Allow-Headers': 'Content-Type'
                }
            )
        else:
            # 非流式输出
            audit_results = {}
            compliant_count = 0
            non_compliant_count = 0
            
            # 对每个检查项进行审核
            for item in filtered_items:
                category = item.get('分类', '未知类别')
                scenario = item.get('专项施工方案严重缺陷情形', '')
                
                # 执行检查
                check_result = auditor.check_category_scenario(category, scenario, top_k=3)
                
                # 简单的合规性判断（这里可以接入LLM进行更智能的判断）
                is_compliant = len(check_result['plan_content']) > 0
                
                if is_compliant:
                    compliant_count += 1
                else:
                    non_compliant_count += 1
                
                audit_results[f"{category}_{item.get('序号', '')}"] = {
                    'category': category,
                    'scenario': scenario,
                    'plan_content': check_result['plan_content'],
                    'check_items': check_result['check_items'],
                    'is_compliant': is_compliant
                }
            
            return jsonify({
                'status': 'success',
                'audit_results': audit_results,
                'summary': {
                    'total_checks': len(filtered_items),
                    'compliant_count': compliant_count,
                    'non_compliant_count': non_compliant_count
                }
            }), 200
        
    except Exception as e:
        logger.error(f"完整审核时发生错误: {str(e)}")
        return jsonify({'status': 'error', 'message': f'完整审核时发生错误: {str(e)}'}), 500

@api_ra_check.route('/ra_check/status', methods=['GET'])
@swag_from(get_status_swagger)
def get_status():
    """查看系统状态"""
    try:
        # 获取文件管理器
        file_manager = FileManager(CACHE_DIR)
        
        # 获取所有文件映射
        all_files = file_manager.get_all_files()
        
        # 获取当前内存中的auditor缓存
        loaded_plans = []
        if hasattr(api_ra_check, 'auditor_cache'):
            for plan_id, auditor in api_ra_check.auditor_cache.items():
                file_info = file_manager.get_file_info(plan_id)
                plan_info = {
                    'plan_id': plan_id,
                    'chunks_count': len(auditor.chunks) if auditor.chunks else 0,
                    'text_length': len(auditor.plan_content),
                    'check_items_count': len(auditor.check_items),
                    'original_filename': file_info.get('original_filename', 'unknown') if file_info else 'unknown',
                    'upload_time': file_info.get('upload_time', '') if file_info else '',
                    'embedding_model': file_info.get('embedding_model', '') if file_info else ''
                }
                loaded_plans.append(plan_info)
        
        # 获取uploads文件夹结构信息
        upload_structure = []
        if os.path.exists(UPLOAD_FOLDER):
            for item in os.listdir(UPLOAD_FOLDER):
                item_path = os.path.join(UPLOAD_FOLDER, item)
                if os.path.isdir(item_path):
                    # 这是一个时间戳文件夹
                    folder_info = {
                        'timestamp_folder': item,
                        'created_time': datetime.fromtimestamp(os.path.getctime(item_path)).isoformat(),
                        'files': []
                    }
                    
                    # 列出文件夹中的文件
                    try:
                        for file in os.listdir(item_path):
                            file_path = os.path.join(item_path, file)
                            if os.path.isfile(file_path):
                                file_size = os.path.getsize(file_path)
                                folder_info['files'].append({
                                    'filename': file,
                                    'size_bytes': file_size,
                                    'size_mb': round(file_size / (1024 * 1024), 2)
                                })
                    except Exception as e:
                        logger.warning(f"读取文件夹 {item} 时出错: {e}")
                    
                    upload_structure.append(folder_info)
        
        return jsonify({
            'status': 'success',
            'loaded_plans': loaded_plans,
            'all_cached_files': all_files,
            'upload_structure': upload_structure,
            'system_info': {
                'cache_dir': CACHE_DIR,
                'upload_folder': UPLOAD_FOLDER,
                'allowed_extensions': list(ALLOWED_EXTENSIONS),
                'total_cached_files': len(all_files),
                'upload_folders_count': len(upload_structure)
            }
        }), 200
        
    except Exception as e:
        logger.error(f"查看状态时发生错误: {str(e)}")
        return jsonify({'status': 'error', 'message': f'查看状态时发生错误: {str(e)}'}), 500

@api_ra_check.route('/ra_check/files', methods=['GET'])
@swag_from(list_files_swagger)
def list_files():
    """列出所有已上传的文件"""
    try:
        file_manager = FileManager(CACHE_DIR)
        all_files = file_manager.get_all_files()
        
        return jsonify({
            'status': 'success',
            'files': all_files,
            'count': len(all_files)
        }), 200
        
    except Exception as e:
        logger.error(f"获取文件列表时发生错误: {str(e)}")
        return jsonify({'status': 'error', 'message': f'获取文件列表时发生错误: {str(e)}'}), 500

@api_ra_check.route('/ra_check/files/<file_hash>', methods=['DELETE'])
@swag_from(delete_file_swagger)
def delete_file(file_hash):
    """删除指定的文件和相关缓存"""
    try:
        file_manager = FileManager(CACHE_DIR)
        
        # 删除文件映射和缓存
        success = file_manager.delete_file_mapping(file_hash)
        
        if not success:
            return jsonify({'status': 'error', 'message': '文件不存在'}), 404
        
        # 从内存缓存中删除
        if hasattr(api_ra_check, 'auditor_cache') and file_hash in api_ra_check.auditor_cache:
            del api_ra_check.auditor_cache[file_hash]
        
        return jsonify({
            'status': 'success',
            'message': '文件删除成功',
            'file_hash': file_hash
        }), 200
        
    except Exception as e:
        logger.error(f"删除文件时发生错误: {str(e)}")
        return jsonify({'status': 'error', 'message': f'删除文件时发生错误: {str(e)}'}), 500

@api_ra_check.route('/ra_check/uploads', methods=['GET'])
@swag_from(list_upload_folders_swagger)
def list_upload_folders():
    """列出所有上传时间戳文件夹"""
    try:
        upload_structure = []
        total_size = 0
        
        if os.path.exists(UPLOAD_FOLDER):
            for item in os.listdir(UPLOAD_FOLDER):
                item_path = os.path.join(UPLOAD_FOLDER, item)
                if os.path.isdir(item_path):
                    folder_size = 0
                    file_count = 0
                    files_info = []
                    
                    # 获取文件夹信息
                    try:
                        for file in os.listdir(item_path):
                            file_path = os.path.join(item_path, file)
                            if os.path.isfile(file_path):
                                file_size = os.path.getsize(file_path)
                                folder_size += file_size
                                file_count += 1
                                files_info.append({
                                    'filename': file,
                                    'size_bytes': file_size,
                                    'size_mb': round(file_size / (1024 * 1024), 2),
                                    'modified_time': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
                                })
                    except Exception as e:
                        logger.warning(f"读取文件夹 {item} 时出错: {e}")
                    
                    folder_info = {
                        'timestamp_folder': item,
                        'created_time': datetime.fromtimestamp(os.path.getctime(item_path)).isoformat(),
                        'file_count': file_count,
                        'total_size_bytes': folder_size,
                        'total_size_mb': round(folder_size / (1024 * 1024), 2),
                        'files': files_info
                    }
                    
                    upload_structure.append(folder_info)
                    total_size += folder_size
        
        # 按时间排序，最新的在前
        upload_structure.sort(key=lambda x: x['created_time'], reverse=True)
        
        return jsonify({
            'status': 'success',
            'upload_folders': upload_structure,
            'summary': {
                'total_folders': len(upload_structure),
                'total_size_bytes': total_size,
                'total_size_mb': round(total_size / (1024 * 1024), 2)
            }
        }), 200
        
    except Exception as e:
        logger.error(f"获取上传文件夹列表时发生错误: {str(e)}")
        return jsonify({'status': 'error', 'message': f'获取上传文件夹列表时发生错误: {str(e)}'}), 500

@api_ra_check.route('/ra_check/uploads/cleanup', methods=['POST'])
@swag_from(cleanup_uploads_swagger)
def cleanup_uploads():
    """清理旧的上传文件夹"""
    try:
        data = request.get_json() or {}
        days_old = data.get('days_old', 7)  # 默认清理7天前的文件夹
        
        # 验证参数
        if not isinstance(days_old, (int, float)) or days_old < 0:
            return jsonify({'status': 'error', 'message': 'days_old 必须是非负数'}), 400
        
        cleaned_count = cleanup_old_upload_folders(days_old)
        
        return jsonify({
            'status': 'success',
            'message': f'清理完成',
            'cleaned_folders_count': cleaned_count,
            'criteria': f'清理了 {days_old} 天前的文件夹'
        }), 200
        
    except Exception as e:
        logger.error(f"清理上传文件夹时发生错误: {str(e)}")
        return jsonify({'status': 'error', 'message': f'清理上传文件夹时发生错误: {str(e)}'}), 500

def load_auditor_from_cache(plan_id):
    """从缓存中加载auditor"""
    try:
        # 检查内存缓存
        if hasattr(api_ra_check, 'auditor_cache') and plan_id in api_ra_check.auditor_cache:
            return api_ra_check.auditor_cache[plan_id]
        
        # 从文件缓存加载
        file_manager = FileManager(CACHE_DIR)
        file_info = file_manager.get_file_info(plan_id)
        
        if not file_info:
            return None
        
        # 构建auditor
        auditor = PlanAuditor(
            plan_content="",  # 暂时为空，后面会从缓存加载
            check_list_file='data/checklist/weakness_list.jsonl',
            embedding_model=file_info.get('embedding_model', 'nomic-embed-text'),
            openai_api_key='ollama',
            openai_api_base='http://59.77.7.24:11434/v1/',
            cache_dir=CACHE_DIR,
            original_filename=file_info.get('original_filename')
        )
        
        # 设置hash
        auditor.file_hash = plan_id
        
        # 从缓存加载嵌入
        cache_files = file_info.get("cache_files", {})
        chunk_file = cache_files.get("chunks")
        emb_file = cache_files.get("embeddings")
        faiss_file = cache_files.get("faiss_index")
        
        if chunk_file and emb_file and faiss_file:
            auditor.load_embeddings(chunk_file, emb_file, faiss_file)
            
            # 读取plan_content
            if os.path.exists(chunk_file):
                with open(chunk_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                    auditor.plan_content = content
        
        # 缓存到内存
        if not hasattr(api_ra_check, 'auditor_cache'):
            api_ra_check.auditor_cache = {}
        api_ra_check.auditor_cache[plan_id] = auditor
        
        return auditor
        
    except Exception as e:
        logger.error(f"从缓存加载auditor失败: {e}")
        return None

@api_ra_check.route('/ra_check/simple_search', methods=['POST'])
@swag_from(simple_search_swagger)
def simple_search():
    """简单检索接口，只返回相关文档片段"""
    try:
        data = request.get_json()
        plan_id = data.get('plan_id')
        query = data.get('query')
        top_k = data.get('top_k', 5)
        
        if not plan_id or not query:
            return jsonify({'status': 'error', 'message': '缺少必要参数'}), 400
        
        # 加载auditor
        auditor = load_auditor_from_cache(plan_id)
        
        if not auditor:
            return jsonify({'status': 'error', 'message': '方案未找到或加载失败'}), 404
        
        # 搜索相似文本块
        similar_chunks = auditor.search_similar_chunks(query, top_k)
        
        # 格式化结果
        results = []
        for i, chunk_result in enumerate(similar_chunks, 1):
            result = {
                'index': i,
                'similarity': chunk_result.get('similarity', 0),
                'text': chunk_result.get('text', ''),
                'preview': (chunk_result.get('text', '')[:200] + "...") if len(chunk_result.get('text', '')) > 200 else chunk_result.get('text', '')
            }
            results.append(result)
        
        # 获取文档信息
        file_manager = FileManager(CACHE_DIR)
        file_info = file_manager.get_file_info(plan_id)
        filename = file_info.get('original_filename', 'unknown') if file_info else 'unknown'
        
        return jsonify({
            'status': 'success',
            'query': query,
            'document': filename,
            'results': results,
            'total_results': len(results)
        }), 200
        
    except Exception as e:
        logger.error(f"简单检索发生错误: {str(e)}")
        return jsonify({'status': 'error', 'message': f'检索发生错误: {str(e)}'}), 500

@api_ra_check.route('/ra_check/stream_query', methods=['POST'])
@swag_from(stream_query_swagger)
def stream_query():
    """基于检索增强的流式查询接口"""
    # 在生成器外部提取请求数据，避免上下文问题
    try:
        data = request.get_json()
        plan_id = data.get('plan_id')
        query = data.get('query')
        top_k = data.get('top_k', 5)
        model_name = data.get('model', 'qwen2.5:7b')  # 默认模型
        
        if not plan_id or not query:
            return jsonify({'status': 'error', 'message': '缺少必要参数'}), 400
    except Exception as e:
        return jsonify({'status': 'error', 'message': f'请求解析错误: {str(e)}'}), 400
    
    def generate_stream_response():
        try:
            # 发送开始信号
            yield f"data: {json.dumps({'type': 'start', 'message': '开始查询...'})}\n\n"
            
            # 加载auditor
            yield f"data: {json.dumps({'type': 'progress', 'message': '加载文档缓存...'})}\n\n"
            auditor = load_auditor_from_cache(plan_id)
            
            if not auditor:
                yield f"data: {json.dumps({'type': 'error', 'message': '方案未找到或加载失败'})}\n\n"
                return
            
            # 获取文档信息
            file_manager = FileManager(CACHE_DIR)
            file_info = file_manager.get_file_info(plan_id)
            
            filename = file_info.get('original_filename', 'unknown') if file_info else 'unknown'
            chunks_count = len(auditor.chunks)
            info_message = f'已加载文档: {filename} ({chunks_count} 个文本块)'
            yield f"data: {json.dumps({'type': 'info', 'message': info_message})}\n\n"
            
            # 执行检索
            yield f"data: {json.dumps({'type': 'progress', 'message': f'正在检索相关内容...'})}\n\n"
            
            # 搜索相似文本块
            similar_chunks = auditor.search_similar_chunks(query, top_k)
            
            yield f"data: {json.dumps({'type': 'progress', 'message': f'找到 {len(similar_chunks)} 个相关文本块，正在生成回复...'})}\n\n"
            
            # 构建检索上下文
            context_texts = []
            for chunk_result in similar_chunks:
                context_texts.append(chunk_result.get('text', ''))
            
            context = "\n\n".join(context_texts)
            
            # 构建prompt
            system_prompt = """你是一个专业的施工方案审核助手。请基于提供的施工方案内容，回答用户的问题。

注意事项：
1. 请严格基于提供的施工方案内容进行回答
2. 如果方案中没有相关信息，请明确说明
3. 回答要专业、准确、具体
4. 可以适当引用方案中的具体条文或数据"""

            user_prompt = f"""施工方案内容：
{context}

用户问题：{query}

请基于上述施工方案内容回答用户问题："""

            # 调用ollama API进行流式生成
            import requests
            
            ollama_url = "http://localhost:11434/api/chat"
            payload = {
                "model": model_name,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                "stream": True
            }
            
            try:
                response = requests.post(ollama_url, json=payload, stream=True, timeout=60)
                response.raise_for_status()
                
                # 开始生成回复信号
                yield f"data: {json.dumps({'type': 'generation_start', 'message': '正在生成回复...'})}\n\n"
                
                # 流式处理响应
                full_response = ""
                for line in response.iter_lines():
                    if line:
                        try:
                            chunk_data = json.loads(line)
                            if 'message' in chunk_data and 'content' in chunk_data['message']:
                                content = chunk_data['message']['content']
                                full_response += content
                                
                                # 流式输出每个token
                                yield f"data: {json.dumps({'type': 'token', 'content': content})}\n\n"
                            
                            # 检查是否完成
                            if chunk_data.get('done', False):
                                # 发送完成信号
                                yield f"data: {json.dumps({'type': 'complete', 'message': '生成完成', 'full_response': full_response, 'context_chunks_count': len(similar_chunks)})}\n\n"
                                break
                                
                        except json.JSONDecodeError:
                            continue
                            
            except requests.exceptions.RequestException as e:
                yield f"data: {json.dumps({'type': 'error', 'message': f'调用模型API失败: {str(e)}'})}\n\n"
                return
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'message': f'生成过程中发生错误: {str(e)}'})}\n\n"
                return
            
        except Exception as e:
            logger.error(f"流式查询发生错误: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'message': f'查询发生错误: {str(e)}'})}\n\n"
    
    return Response(
        generate_stream_response(),
        mimetype='text/plain',
        headers={
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive',
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Headers': 'Content-Type'
        }
    )

@api_ra_check.route('/ra_check/available_embeddings', methods=['GET'])
@swag_from(available_embeddings_swagger)
def get_available_embeddings():
    """获取可用的嵌入文件列表"""
    try:
        file_manager = FileManager(CACHE_DIR)
        all_files = file_manager.get_all_files()
        
        available_embeddings = []
        for file_info in all_files:
            file_hash = file_info.get('file_hash')
            embedding_info = {
                'plan_id': file_hash,
                'original_filename': file_info.get('original_filename', 'unknown'),
                'upload_time': file_info.get('upload_time', ''),
                'embedding_model': file_info.get('embedding_model', ''),
                'chunks_count': file_info.get('chunks_count', 0),
                'text_length': file_info.get('text_length', 0),
                'is_loaded': False
            }
            
            # 检查是否已在内存中加载
            if hasattr(api_ra_check, 'auditor_cache') and file_hash in api_ra_check.auditor_cache:
                embedding_info['is_loaded'] = True
            
            available_embeddings.append(embedding_info)
        
        # 按上传时间排序
        available_embeddings.sort(key=lambda x: x['upload_time'], reverse=True)
        
        return jsonify({
            'status': 'success',
            'available_embeddings': available_embeddings,
            'count': len(available_embeddings)
        }), 200
        
    except Exception as e:
        logger.error(f"获取可用嵌入文件列表时发生错误: {str(e)}")
        return jsonify({'status': 'error', 'message': f'获取可用嵌入文件列表时发生错误: {str(e)}'}), 500

# 新增：批量检查API - 支持上传检查项JSON和文档
@api_ra_check.route('/ra_check/batch_check', methods=['POST'])
@swag_from(batch_check_swagger)
def batch_check():
    """
    批量检查接口：上传检查项JSON文件和待检查文档，逐个返回检查结果
    
    支持的文件格式：
    - checklist: JSON文件（类似weakness_list.jsonl格式）
    - document: DOCX文件（待检查的文档）
    
    返回格式：
    - 每个检查项的原文依据、判断结果和概率
    """
    try:
        # 检查是否有文件上传
        if 'checklist' not in request.files or 'document' not in request.files:
            return jsonify({
                'status': 'error', 
                'message': '请同时上传检查项文件(checklist)和待检查文档(document)'
            }), 400
        
        checklist_file = request.files['checklist']
        document_file = request.files['document']
        
        # 验证文件名
        if checklist_file.filename == '' or document_file.filename == '':
            return jsonify({'status': 'error', 'message': '未选择文件'}), 400
        
        # 验证文件类型
        if not checklist_file.filename.lower().endswith(('.json', '.jsonl')):
            return jsonify({'status': 'error', 'message': '检查项文件必须是JSON或JSONL格式'}), 400
        
        if not document_file.filename.lower().endswith('.docx'):
            return jsonify({'status': 'error', 'message': '待检查文档必须是DOCX格式'}), 400
        
        # 创建临时文件夹保存上传的文件
        upload_subfolder, timestamp_folder = ensure_upload_subfolder()
        
        # 保存检查项文件
        checklist_filename = secure_filename(checklist_file.filename)
        checklist_path = os.path.join(upload_subfolder, checklist_filename)
        checklist_file.save(checklist_path)
        
        # 保存文档文件
        document_filename = secure_filename(document_file.filename)
        document_path = os.path.join(upload_subfolder, document_filename)
        document_file.save(document_path)
        
        # 解析检查项文件
        checklist_items = []
        try:
            with open(checklist_path, 'r', encoding='utf-8') as f:
                if checklist_filename.endswith('.jsonl'):
                    # JSONL格式，每行一个JSON对象
                    for line in f:
                        line = line.strip()
                        if line:
                            checklist_items.append(json.loads(line))
                else:
                    # JSON格式，可能是数组或单个对象
                    content = json.load(f)
                    if isinstance(content, list):
                        checklist_items = content
                    else:
                        checklist_items = [content]
        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'解析检查项文件失败: {str(e)}'
            }), 400
        
        if not checklist_items:
            return jsonify({
                'status': 'error',
                'message': '检查项文件为空或格式不正确'
            }), 400
        
        # 提取文档内容
        try:
            document_content = extract_text_from_docx(document_path)
        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'提取文档内容失败: {str(e)}'
            }), 400
        
        # 获取配置参数
        embedding_model = request.form.get('embedding_model', 'nomic-embed-text:latest')
        openai_api_key = request.form.get('openai_api_key', 'ollama')
        openai_api_base = request.form.get('openai_api_base', 'http://59.77.7.24:11434/v1/')
        top_k = int(request.form.get('top_k', 5))
        chat_model = request.form.get('chat_model', 'qwen2.5:32b')  # 提前提取chat_model参数
        stream = request.form.get('stream', 'false').lower() == 'true'
        
        # 创建临时检查项文件（JSONL格式）
        temp_checklist_path = os.path.join(upload_subfolder, 'temp_checklist.jsonl')
        with open(temp_checklist_path, 'w', encoding='utf-8') as f:
            for item in checklist_items:
                f.write(json.dumps(item, ensure_ascii=False) + '\n')
        
        # 初始化审查器
        try:
            auditor = PlanAuditor(
                plan_content=document_content,
                check_list_file=temp_checklist_path,
                embedding_model=embedding_model,
                openai_api_key=openai_api_key,
                openai_api_base=openai_api_base,
                cache_dir=CACHE_DIR,
                original_filename=document_filename
            )
            
            # 构建嵌入
            plan_id = auditor.build_or_load_embeddings()
            
        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'初始化审查器失败: {str(e)}'
            }), 500
        
        if stream:
            # 流式输出
            def generate_stream_response():
                try:
                    yield f"data: {json.dumps({'type': 'start', 'message': f'开始批量检查，共{len(checklist_items)}个检查项...'})}\n\n"
                    
                    check_results = []
                    
                    for i, check_item in enumerate(checklist_items, 1):
                        try:
                            logger.info(f"正在检查第 {i}/{len(checklist_items)} 项: {check_item}")
                            
                            # 获取检查项信息
                            category = check_item.get('分类', '未知分类')
                            check_scenario = check_item.get('专项施工方案严重缺陷情形', '')
                            item_number = check_item.get('序号', str(i))
                            
                            if not check_scenario:
                                result = {
                                    'item_number': item_number,
                                    'category': category,
                                    'check_scenario': '检查项信息不完整',
                                    'evidence': '',
                                    'judgment': '无法判断',
                                    'probability': 0.0,
                                    'error': '检查项缺少"专项施工方案严重缺陷情形"字段'
                                }
                                check_results.append(result)
                                yield f"data: {json.dumps({'type': 'item_complete', 'item_index': i, 'result': result})}\n\n"
                                continue
                            
                            # 搜索相关文本片段
                            similar_chunks = auditor.search_similar_chunks(check_scenario, top_k=top_k)
                            
                            # 组合检索到的文本作为证据
                            evidence_texts = []
                            for chunk_dict in similar_chunks:
                                chunk_text = chunk_dict.get("text", "")
                                similarity = chunk_dict.get("similarity", 0.0)
                                evidence_texts.append(f"相关度{similarity:.3f}: {chunk_text}")
                            
                            evidence = "\n".join(evidence_texts)
                            
                            # 使用大模型进行智能判断
                            if similar_chunks:
                                # 构建上下文内容
                                context = "\n".join([chunk_dict.get("text", "") for chunk_dict in similar_chunks])
                                
                                                                # 构建判断prompt
                                messages = [
                                    {
                                        "role": "system",
                                        "content": CONSTRUCTION_EXPERT_SYSTEM
                                    },
                                    {
                                        "role": "user",
                                        "content": get_batch_check_prompt(check_scenario, category, context)
                                    }
                                ]
                                
                                # 流式生成分析
                                yield f"data: {json.dumps({'type': 'generation_start', 'message': f'正在分析第{i}项...', 'item_index': i})}\n\n"
                                
                                llm_response = ""
                                for token in auditor.embedder.generate_text_stream(messages, model=chat_model):
                                    llm_response += token
                                    yield f"data: {json.dumps({'type': 'token', 'content': token, 'item_index': i, 'category': category})}\n\n"
                                
                                # 解析大模型回复
                                judgment = parse_llm_judgment(llm_response)
                                probability = parse_confidence_score(llm_response)
                                
                                check_result = llm_response
                                
                            else:
                                # 没有找到相关内容
                                judgment = "无法判断"
                                probability = 0.0
                                check_result = "未找到相关文档内容"
                            
                            result = {
                                'item_number': item_number,
                                'category': category,
                                'check_scenario': check_scenario,
                                'evidence': evidence,
                                'judgment': judgment,
                                'probability': round(probability, 3),
                                'detailed_result': check_result,
                                'chunk_count': len(similar_chunks)
                            }
                            check_results.append(result)
                            
                            yield f"data: {json.dumps({'type': 'item_complete', 'item_index': i, 'result': result})}\n\n"
                            
                        except Exception as e:
                            logger.error(f"检查第 {i} 项时发生错误: {str(e)}")
                            result = {
                                'item_number': check_item.get('序号', str(i)),
                                'category': check_item.get('分类', '未知分类'),
                                'check_scenario': check_item.get('专项施工方案严重缺陷情形', ''),
                                'evidence': '',
                                'judgment': '检查失败',
                                'probability': 0.0,
                                'error': str(e)
                            }
                            check_results.append(result)
                            yield f"data: {json.dumps({'type': 'item_error', 'item_index': i, 'result': result})}\n\n"
                    
                    # 计算总体统计
                    total_items = len(check_results)
                    compliant_items = len([r for r in check_results if r['judgment'] == '合规'])
                    non_compliant_items = len([r for r in check_results if r['judgment'] == '不合规'])
                    failed_items = len([r for r in check_results if r['judgment'] == '检查失败'])
                    
                    data = {
                        'upload_folder': timestamp_folder,
                        'document_filename': document_filename,
                        'checklist_filename': checklist_filename,
                        'plan_id': plan_id,
                        'summary': {
                            'total_items': total_items,
                            'compliant_items': compliant_items,
                            'non_compliant_items': non_compliant_items,
                            'failed_items': failed_items,
                            'compliance_rate': round((compliant_items / total_items * 100), 2) if total_items > 0 else 0
                        },
                        'check_results': check_results
                    }
                    
                    yield f"data: {json.dumps({'type': 'complete', 'message': '批量检查完成', 'data': data})}\n\n"
                    
                except Exception as e:
                    logger.error(f"流式批量检查发生错误: {str(e)}")
                    yield f"data: {json.dumps({'type': 'error', 'message': f'检查发生错误: {str(e)}'})}\n\n"
            
            return Response(
                generate_stream_response(),
                mimetype='text/plain',
                headers={
                    'Cache-Control': 'no-cache',
                    'Connection': 'keep-alive',
                    'Access-Control-Allow-Origin': '*',
                    'Access-Control-Allow-Headers': 'Content-Type'
                }
            )
        else:
            # 非流式输出
            # 逐个检查每个检查项
            check_results = []
            
            for i, check_item in enumerate(checklist_items, 1):
                try:
                    logger.info(f"正在检查第 {i}/{len(checklist_items)} 项: {check_item}")
                    
                    # 获取检查项信息
                    category = check_item.get('分类', '未知分类')
                    check_scenario = check_item.get('专项施工方案严重缺陷情形', '')
                    item_number = check_item.get('序号', str(i))
                    
                    if not check_scenario:
                        check_results.append({
                            'item_number': item_number,
                            'category': category,
                            'check_scenario': '检查项信息不完整',
                            'evidence': '',
                            'judgment': '无法判断',
                            'probability': 0.0,
                            'error': '检查项缺少"专项施工方案严重缺陷情形"字段'
                        })
                        continue
                    
                    # 搜索相关文本片段
                    similar_chunks = auditor.search_similar_chunks(check_scenario, top_k=top_k)
                    
                    # 组合检索到的文本作为证据
                    evidence_texts = []
                    for chunk_dict in similar_chunks:
                        chunk_text = chunk_dict.get("text", "")
                        similarity = chunk_dict.get("similarity", 0.0)
                        evidence_texts.append(f"相关度{similarity:.3f}: {chunk_text}")
                    
                    evidence = "\n".join(evidence_texts)
                    
                    # 使用大模型进行智能判断
                    if similar_chunks:
                        # 构建上下文内容
                        context = "\n".join([chunk_dict.get("text", "") for chunk_dict in similar_chunks])
                        
                        # 构建判断prompt
                        messages = [
                            {
                                "role": "system",
                                "content": CONSTRUCTION_EXPERT_SYSTEM
                            },
                            {
                                "role": "user",
                                "content": get_batch_check_prompt(check_scenario, category, context)
                            }
                        ]
                        
                        # 调用大模型进行判断
                        try:
                            llm_response = auditor.embedder.generate_text(
                                messages=messages,
                                model=chat_model,
                                temperature=0.1
                            )
                            
                            # 解析大模型回复
                            judgment = parse_llm_judgment(llm_response)
                            probability = parse_confidence_score(llm_response)
                            detailed_analysis = llm_response
                            
                            check_result = llm_response
                            
                        except Exception as e:
                            logger.error(f"大模型调用失败: {str(e)}")
                            # 降级到简单判断
                            judgment = "合规"  # 默认合规
                            probability = 0.5
                            check_result = f"大模型调用失败，降级判断: {str(e)}"
                            detailed_analysis = "大模型调用失败"
                            
                    else:
                        # 没有找到相关内容
                        judgment = "无法判断"
                        probability = 0.0
                        check_result = "未找到相关文档内容"
                        detailed_analysis = "未找到相关文档内容"
                    
                    check_results.append({
                        'item_number': item_number,
                        'category': category,
                        'check_scenario': check_scenario,
                        'evidence': evidence,
                        'judgment': judgment,
                        'probability': round(probability, 3),
                        'detailed_result': check_result,
                        'chunk_count': len(similar_chunks)
                    })
                    
                except Exception as e:
                    logger.error(f"检查第 {i} 项时发生错误: {str(e)}")
                    check_results.append({
                        'item_number': check_item.get('序号', str(i)),
                        'category': check_item.get('分类', '未知分类'),
                        'check_scenario': check_item.get('专项施工方案严重缺陷情形', ''),
                        'evidence': '',
                        'judgment': '检查失败',
                        'probability': 0.0,
                        'error': str(e)
                    })
            
            # 计算总体统计
            total_items = len(check_results)
            compliant_items = len([r for r in check_results if r['judgment'] == '合规'])
            non_compliant_items = len([r for r in check_results if r['judgment'] == '不合规'])
            failed_items = len([r for r in check_results if r['judgment'] == '检查失败'])
            
            return jsonify({
                'status': 'success',
                'message': '批量检查完成',
                'data': {
                    'upload_folder': timestamp_folder,
                    'document_filename': document_filename,
                    'checklist_filename': checklist_filename,
                    'plan_id': plan_id,
                    'summary': {
                        'total_items': total_items,
                        'compliant_items': compliant_items,
                        'non_compliant_items': non_compliant_items,
                        'failed_items': failed_items,
                        'compliance_rate': round(compliant_items / total_items * 100, 2) if total_items > 0 else 0
                    },
                    'check_results': check_results
                }
            })
            
    except Exception as e:
        logger.error(f"批量检查API发生错误: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'服务器内部错误: {str(e)}'
        }), 500

@api_ra_check.route('/ra_check/cite_check', methods=['POST'])
@swag_from(cite_check_swagger)
def cite_check():
    """
    引用检查接口：上传引用检查文件和待检查文档，检查文档是否正确引用了指定的条目
    
    支持的文件格式：
    - cite_list: JSON或JSONL文件（包含引用条目信息）
    - document: DOCX文件（待检查的文档）
    
    返回格式：
    - 每个引用条目的检查结果、引用状态和准确性评分
    """
    try:
        # 检查是否有文件上传
        if 'cite_list' not in request.files or 'document' not in request.files:
            return jsonify({
                'status': 'error', 
                'message': '请同时上传引用检查文件(cite_list)和待检查文档(document)'
            }), 400
        
        cite_list_file = request.files['cite_list']
        document_file = request.files['document']
        
        # 验证文件名
        if cite_list_file.filename == '' or document_file.filename == '':
            return jsonify({'status': 'error', 'message': '未选择文件'}), 400
        
        # 验证文件类型
        if not cite_list_file.filename.lower().endswith(('.json', '.jsonl')):
            return jsonify({'status': 'error', 'message': '引用检查文件必须是JSON或JSONL格式'}), 400
        
        if not document_file.filename.lower().endswith('.docx'):
            return jsonify({'status': 'error', 'message': '待检查文档必须是DOCX格式'}), 400
        
        # 创建临时文件夹保存上传的文件
        upload_subfolder, timestamp_folder = ensure_upload_subfolder()
        
        # 保存引用检查文件
        cite_list_filename = secure_filename(cite_list_file.filename)
        cite_list_path = os.path.join(upload_subfolder, cite_list_filename)
        cite_list_file.save(cite_list_path)
        
        # 保存文档文件
        document_filename = secure_filename(document_file.filename)
        document_path = os.path.join(upload_subfolder, document_filename)
        document_file.save(document_path)
        
        # 解析引用检查文件
        citation_items = []
        try:
            with open(cite_list_path, 'r', encoding='utf-8') as f:
                if cite_list_filename.endswith('.jsonl'):
                    # JSONL格式，每行一个JSON对象
                    for line in f:
                        line = line.strip()
                        if line:
                            citation_items.append(json.loads(line))
                else:
                    # JSON格式，可能是数组或单个对象
                    content = json.load(f)
                    if isinstance(content, list):
                        citation_items = content
                    else:
                        citation_items = [content]
        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'解析引用检查文件失败: {str(e)}'
            }), 400
        
        if not citation_items:
            return jsonify({
                'status': 'error',
                'message': '引用检查文件为空或格式不正确'
            }), 400
        
        # 提取文档内容
        try:
            document_content = extract_text_from_docx(document_path)
        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'提取文档内容失败: {str(e)}'
            }), 400
        
        # 获取配置参数
        embedding_model = request.form.get('embedding_model', 'nomic-embed-text:latest')
        openai_api_key = request.form.get('openai_api_key', 'ollama')
        openai_api_base = request.form.get('openai_api_base', 'http://59.77.7.24:11434/v1/')
        top_k = int(request.form.get('top_k', 5))
        stream = request.form.get('stream', 'false').lower() == 'true'
        
        # 创建临时引用检查文件（JSONL格式）
        temp_cite_list_path = os.path.join(upload_subfolder, 'temp_cite_list.jsonl')
        with open(temp_cite_list_path, 'w', encoding='utf-8') as f:
            for item in citation_items:
                f.write(json.dumps(item, ensure_ascii=False) + '\n')
        
        # 初始化审查器
        try:
            auditor = PlanAuditor(
                plan_content=document_content,
                check_list_file=temp_cite_list_path,
                embedding_model=embedding_model,
                openai_api_key=openai_api_key,
                openai_api_base=openai_api_base,
                cache_dir=CACHE_DIR,
                original_filename=document_filename
            )
            
            # 构建嵌入
            plan_id = auditor.build_or_load_embeddings()
            
        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'初始化审查器失败: {str(e)}'
            }), 500
        
        # 逐个检查每个引用条目
        citation_results = []
        
        for i, citation_item in enumerate(citation_items, 1):
            try:
                logger.info(f"正在检查第 {i}/{len(citation_items)} 个引用条目: {citation_item}")
                
                # 获取引用条目信息（支持多种字段格式）
                # 学术文献格式
                citation_id = citation_item.get('id', citation_item.get('citation_id', str(i)))
                title = citation_item.get('title', citation_item.get('标题', ''))
                authors = citation_item.get('authors', citation_item.get('作者', ''))
                publication = citation_item.get('publication', citation_item.get('出版物', ''))
                year = citation_item.get('year', citation_item.get('年份', ''))
                citation_text = citation_item.get('citation_text', citation_item.get('引用文本', ''))
                
                # 标准规范格式（优先使用标准规范字段）
                standard_code = citation_item.get('标准编号', '')
                standard_name = citation_item.get('标准名称', '')
                issuing_dept = citation_item.get('发布部门', '')
                implementation_date = citation_item.get('实施日期', '')
                status = citation_item.get('状态', '')
                
                # 如果是标准规范格式，重新组织字段
                if standard_code or standard_name:
                    citation_id = citation_id or standard_code
                    title = title or standard_name
                    authors = authors or issuing_dept
                    publication = publication or '标准规范'
                    if implementation_date and not year:
                        # 从实施日期提取年份
                        try:
                            year = implementation_date.split('-')[0]
                        except:
                            year = ''
                
                # 如果没有明确的引用文本，则从其他字段构建查询
                if not citation_text:
                    search_components = []
                    if standard_code:
                        search_components.append(standard_code)
                    if standard_name:
                        search_components.append(standard_name)
                    elif title:
                        search_components.append(title)
                    if issuing_dept:
                        search_components.append(issuing_dept)
                    elif authors:
                        search_components.append(authors)
                    if implementation_date:
                        search_components.append(implementation_date)
                    elif year:
                        search_components.append(str(year))
                    
                    citation_text = ' '.join(search_components)
                
                if not citation_text.strip():
                    citation_results.append({
                        'citation_id': citation_id,
                        'title': title,
                        'authors': authors,
                        'publication': publication,
                        'year': year,
                        'standard_code': standard_code,
                        'standard_name': standard_name,
                        'issuing_dept': issuing_dept,
                        'implementation_date': implementation_date,
                        'status': status,
                        'citation_text': citation_text,
                        'evidence': '',
                        'citation_status': '检查失败',
                        'accuracy_score': 0.0,
                        'detailed_result': '引用条目信息不完整，无法进行检查',
                        'chunk_count': 0,
                        'error': '缺少可检索的引用文本信息'
                    })
                    continue
                
                # 搜索相关文本片段
                similar_chunks = auditor.search_similar_chunks(citation_text, top_k=top_k)
                
                # 组合检索到的文本作为证据
                evidence_texts = []
                for chunk_dict in similar_chunks:
                    chunk_text = chunk_dict.get("text", "")
                    similarity = chunk_dict.get("similarity", 0.0)
                    evidence_texts.append(f"相关度{similarity:.3f}: {chunk_text}")
                
                evidence = "\n".join(evidence_texts)
                
                # 使用大模型进行引用检查
                if similar_chunks:
                    # 构建上下文内容
                    context = "\n".join([chunk_dict.get("text", "") for chunk_dict in similar_chunks])
                    
                    # 构建引用检查prompt（针对标准规范或学术文献）
                    if standard_code or standard_name:
                        # 标准规范检查prompt
                        messages = [
                            {
                                "role": "system",
                                "content": "你是一位专业的工程技术文档审查专家。请基于提供的文档内容，判断文档是否正确引用了指定的技术标准规范。"
                            },
                            {
                                "role": "user",
                                "content": f"""
请分析以下技术文档内容，判断是否正确引用了指定的技术标准规范。

【待检查的标准规范】:
- 标准编号: {standard_code}
- 标准名称: {standard_name}
- 发布部门: {issuing_dept}
- 实施日期: {implementation_date}
- 状态: {status}
- 查询文本: {citation_text}

【文档相关内容】:
{context}

请按以下格式回答：
1. 引用状态：[正确引用/缺失引用/引用有误/引用不完整]
2. 准确性评分：[0.0-1.0之间的数值，表示引用的准确性]
3. 分析说明：[详细说明分析过程和判断依据]

判断标准：
- 正确引用：文档中明确提及了该标准的编号、名称等关键信息
- 缺失引用：文档中应该引用该标准但完全没有提及
- 引用有误：文档中提及了该标准但信息有错误（编号错误、名称错误等）
- 引用不完整：文档中提及了该标准但引用信息不完整（如只有编号没有名称）
- 准确性评分反映标准引用的完整性和准确性
"""
                            }
                        ]
                    else:
                        # 学术文献检查prompt
                        messages = [
                            {
                                "role": "system",
                                "content": "你是一位专业的学术引用检查专家。请基于提供的文档内容，判断文档是否正确引用了指定的学术条目。"
                            },
                            {
                                "role": "user",
                                "content": f"""
请分析以下文档内容，判断是否正确引用了指定的学术条目。

【待检查的引用条目】:
- 标题: {title}
- 作者: {authors}
- 出版物: {publication}
- 年份: {year}
- 查询文本: {citation_text}

【文档相关内容】:
{context}

请按以下格式回答：
1. 引用状态：[正确引用/缺失引用/引用有误/引用不完整]
2. 准确性评分：[0.0-1.0之间的数值，表示引用的准确性]
3. 分析说明：[详细说明分析过程和判断依据]

判断标准：
- 正确引用：文档中包含了该条目的准确引用信息（作者、标题、年份等）
- 缺失引用：文档中应该引用但完全没有提及该条目
- 引用有误：文档中提及了该条目但信息有错误（作者名错误、年份错误等）
- 引用不完整：文档中提及了该条目但引用信息不完整
- 准确性评分反映引用信息的完整性和准确性
"""
                            }
                        ]
                    
                    # 调用大模型进行判断
                    try:
                        llm_response = auditor.embedder.generate_text(
                            messages=messages,
                            model=request.form.get('chat_model', 'qwen2.5:32b'),
                            temperature=0.1
                        )
                        
                        # 解析大模型回复
                        citation_status = "缺失引用"  # 默认状态
                        accuracy_score = 0.0  # 默认评分
                        detailed_analysis = llm_response
                        
                        # 简单的文本解析来提取引用状态
                        response_lower = llm_response.lower()
                        if "正确引用" in response_lower:
                            citation_status = "正确引用"
                        elif "引用有误" in response_lower:
                            citation_status = "引用有误"
                        elif "引用不完整" in response_lower:
                            citation_status = "引用不完整"
                        elif "缺失引用" in response_lower:
                            citation_status = "缺失引用"
                        
                        # 尝试提取准确性评分
                        import re
                        # 尝试多种评分表达模式
                        patterns = [
                            r'准确性评分[：:]\s*([0-9.]+)',
                            r'评分[：:]\s*([0-9.]+)',
                            r'2\.\s*准确性评分[：:]\s*([0-9.]+)'
                        ]
                        
                        for pattern in patterns:
                            score_match = re.search(pattern, llm_response)
                            if score_match:
                                try:
                                    accuracy_score = float(score_match.group(1))
                                    accuracy_score = max(0.0, min(1.0, accuracy_score))  # 确保在0-1范围内
                                    break
                                except ValueError:
                                    continue
                        
                        detailed_result = llm_response
                        
                    except Exception as e:
                        logger.error(f"大模型调用失败: {str(e)}")
                        # 降级到简单判断
                        citation_status = "检查失败"
                        accuracy_score = 0.0
                        detailed_result = f"大模型调用失败，无法完成引用检查: {str(e)}"
                        detailed_analysis = "大模型调用失败"
                        
                    citation_results.append({
                        'citation_id': citation_id,
                        'title': title,
                        'authors': authors,
                        'publication': publication,
                        'year': year,
                        'standard_code': standard_code,
                        'standard_name': standard_name,
                        'issuing_dept': issuing_dept,
                        'implementation_date': implementation_date,
                        'status': status,
                        'citation_text': citation_text,
                        'evidence': evidence,
                        'citation_status': citation_status,
                        'accuracy_score': round(accuracy_score, 3),
                        'detailed_result': detailed_result,
                        'chunk_count': len(similar_chunks)
                    })
                
            except Exception as e:
                logger.error(f"检查第 {i} 个引用条目时发生错误: {str(e)}")
                citation_results.append({
                    'citation_id': citation_item.get('id', citation_item.get('citation_id', citation_item.get('标准编号', str(i)))),
                    'title': citation_item.get('title', citation_item.get('标题', citation_item.get('标准名称', ''))),
                    'authors': citation_item.get('authors', citation_item.get('作者', citation_item.get('发布部门', ''))),
                    'publication': citation_item.get('publication', citation_item.get('出版物', '标准规范')),
                    'year': citation_item.get('year', citation_item.get('年份', '')),
                    'standard_code': citation_item.get('标准编号', ''),
                    'standard_name': citation_item.get('标准名称', ''),
                    'issuing_dept': citation_item.get('发布部门', ''),
                    'implementation_date': citation_item.get('实施日期', ''),
                    'status': citation_item.get('状态', ''),
                    'citation_text': citation_item.get('citation_text', citation_item.get('引用文本', '')),
                    'evidence': '',
                    'citation_status': '检查失败',
                    'accuracy_score': 0.0,
                    'detailed_result': f'检查过程中发生错误: {str(e)}',
                    'chunk_count': 0,
                    'error': str(e)
                })
        
        # 计算总体统计
        total_citations = len(citation_results)
        properly_cited = len([r for r in citation_results if r['citation_status'] == '正确引用'])
        missing_citations = len([r for r in citation_results if r['citation_status'] == '缺失引用'])
        incorrectly_cited = len([r for r in citation_results if r['citation_status'] in ['引用有误', '引用不完整']])
        failed_checks = len([r for r in citation_results if r['citation_status'] == '检查失败'])
        
        return jsonify({
            'status': 'success',
            'message': '引用检查完成',
            'data': {
                'upload_folder': timestamp_folder,
                'document_filename': document_filename,
                'cite_list_filename': cite_list_filename,
                'plan_id': plan_id,
                'summary': {
                    'total_citations': total_citations,
                    'properly_cited': properly_cited,
                    'missing_citations': missing_citations,
                    'incorrectly_cited': incorrectly_cited,
                    'failed_checks': failed_checks,
                    'citation_rate': round(properly_cited / total_citations * 100, 2) if total_citations > 0 else 0
                },
                'citation_results': citation_results
            }
        })
        
    except Exception as e:
        logger.error(f"引用检查API发生错误: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'服务器内部错误: {str(e)}'
        }), 500

@api_ra_check.route('/ra_check/structure_check', methods=['POST'])
@swag_from(structure_check_swagger)
def structure_check():
    """
    文档结构完整性检查接口
    
    支持两种检查模式：
    1. item_by_item: 逐条检查，每个目录项单独分析
    2. chapter_by_chapter: 逐章节检查，同一章节的内容一起分析
    """
    try:
        # 获取上传的文件
        if 'toc_list' not in request.files or 'document' not in request.files:
            return jsonify({
                'status': 'error',
                'message': '请上传目录结构清单文件(toc_list)和待检查文档(document)'
            }), 400
        
        toc_list_file = request.files['toc_list']
        document_file = request.files['document']
        
        if toc_list_file.filename == '' or document_file.filename == '':
            return jsonify({
                'status': 'error',
                'message': '请选择有效的文件'
            }), 400
        
        # 获取参数
        check_mode = request.form.get('check_mode', 'item_by_item')  # 默认逐条检查
        embedding_model = request.form.get('embedding_model', 'nomic-embed-text:latest')
        chat_model = request.form.get('chat_model', 'qwen2.5:32b')
        top_k = int(request.form.get('top_k', 5))
        openai_api_key = request.form.get('openai_api_key', 'ollama')
        openai_api_base = request.form.get('openai_api_base', 'http://59.77.7.24:11434/v1/')
        stream = request.form.get('stream', 'false').lower() == 'true'
        
        # 验证检查模式
        if check_mode not in ['item_by_item', 'chapter_by_chapter']:
            return jsonify({
                'status': 'error',
                'message': '检查模式必须是 item_by_item 或 chapter_by_chapter'
            }), 400
        
        # 创建时间戳文件夹
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]
        upload_folder = os.path.join(UPLOAD_FOLDER, timestamp)
        os.makedirs(upload_folder, exist_ok=True)
        
        # 保存文件
        toc_list_filename = secure_filename(toc_list_file.filename)
        document_filename = secure_filename(document_file.filename)
        
        toc_list_path = os.path.join(upload_folder, toc_list_filename)
        document_path = os.path.join(upload_folder, document_filename)
        
        toc_list_file.save(toc_list_path)
        document_file.save(document_path)
        
        # 验证文件格式
        if not document_filename.lower().endswith('.docx'):
            return jsonify({
                'status': 'error',
                'message': '文档必须是DOCX格式'
            }), 400
        
        if not toc_list_filename.lower().endswith(('.json', '.jsonl')):
            return jsonify({
                'status': 'error',
                'message': '目录结构清单必须是JSON或JSONL格式'
            }), 400
        
        # 解析目录结构清单
        toc_items = []
        try:
            with open(toc_list_path, 'r', encoding='utf-8') as f:
                if toc_list_filename.lower().endswith('.json'):
                    # JSON格式
                    toc_items = json.load(f)
                else:
                    # JSONL格式
                    for line in f:
                        line = line.strip()
                        if line:
                            toc_items.append(json.loads(line))
        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'解析目录结构清单失败: {str(e)}'
            }), 400
        
        if not toc_items:
            return jsonify({
                'status': 'error',
                'message': '目录结构清单为空'
            }), 400
        
        # 处理文档并生成嵌入向量（复用现有逻辑）
        try:
            # 提取文档文本
            doc_text = extract_text_from_docx(document_path)
            if not doc_text:
                return jsonify({
                    'status': 'error',
                    'message': '无法从DOCX文档中提取文本内容'
                }), 400
            
            # 创建临时检查清单文件（PlanAuditor需要，但我们不使用其检查功能）
            temp_check_list = os.path.join(upload_folder, 'temp_check_list.jsonl')
            with open(temp_check_list, 'w', encoding='utf-8') as f:
                f.write('{"temp": "temp"}\n')
            
            # 创建文档处理器
            auditor = PlanAuditor(
                plan_content=doc_text,
                check_list_file=temp_check_list,
                embedding_model=embedding_model,
                openai_api_key=openai_api_key,
                openai_api_base=openai_api_base,
                cache_dir=CACHE_DIR,
                original_filename=document_filename
            )
            
            # 构建嵌入向量
            plan_id = auditor.build_or_load_embeddings()
            
        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'文档处理失败: {str(e)}'
            }), 500
        
        # 根据检查模式进行结构完整性检查
        if check_mode == 'item_by_item':
            check_results = perform_item_by_item_structure_check(toc_items, auditor, chat_model, top_k)
        else:  # chapter_by_chapter
            check_results = perform_chapter_by_chapter_structure_check(toc_items, auditor, chat_model, top_k)
        
        # 计算统计信息
        total_items = len(check_results)
        complete_items = len([r for r in check_results if r['completeness_status'] == '完整'])
        missing_items = len([r for r in check_results if r['completeness_status'] == '缺失'])
        partial_items = len([r for r in check_results if r['completeness_status'] == '部分完整'])
        failed_checks = len([r for r in check_results if r['completeness_status'] == '检查失败'])
        
        completeness_rate = (complete_items / total_items * 100) if total_items > 0 else 0
        
        return jsonify({
            'status': 'success',
            'message': '文档结构完整性检查完成',
            'data': {
                'summary': {
                    'total_items': total_items,
                    'complete_items': complete_items,
                    'missing_items': missing_items,
                    'partial_items': partial_items,
                    'failed_checks': failed_checks,
                    'completeness_rate': round(completeness_rate, 2)
                },
                'check_results': check_results,
                'check_mode': check_mode,
                'toc_list_filename': toc_list_filename,
                'document_filename': document_filename,
                'plan_id': plan_id,
                'upload_folder': timestamp
            }
        })
        
    except Exception as e:
        logger.error(f"结构检查过程中发生错误: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'检查失败: {str(e)}'
        }), 500

def perform_item_by_item_structure_check(toc_items, auditor, chat_model, top_k):
    """逐条检查模式"""
    results = []
    
    for i, item in enumerate(toc_items):
        try:
            chapter = item.get('章节', '')
            name = item.get('名称', '')
            required = item.get('必有', '否')
            item_type = item.get('类型', '')
            ai_applicable = item.get('AI适用', '否')
            description = item.get('说明', '')
            
            # 构建检索查询
            search_query = f"{chapter} {name}"
            
            # 检索相关文档片段
            try:
                similar_chunks_results = auditor.search_similar_chunks(search_query, top_k=top_k)
                
                # 过滤相似度低的结果
                similar_chunks = []
                for result in similar_chunks_results:
                    if result['similarity'] > 0.1:
                        similar_chunks.append((result['text'], result['similarity']))
                
                # 构建证据文本
                evidence_text = ""
                if similar_chunks:
                    evidence_chunks = []
                    for chunk, similarity in similar_chunks:
                        evidence_chunks.append(f"相关度{similarity:.3f}: {chunk}")
                    evidence_text = '\n'.join(evidence_chunks[:3])  # 最多3个相关片段
                
                # 使用AI分析结构完整性
                if ai_applicable == '是' and evidence_text:
                    analysis_result = analyze_structure_completeness_single(
                        item, evidence_text, chat_model, auditor
                    )
                    analysis_result['item_id'] = str(i + 1)
                    results.append(analysis_result)
                else:
                    # 对于不适用AI的项目，进行简单的关键词匹配
                    simple_result = simple_structure_check_single(item, evidence_text, i + 1)
                    results.append(simple_result)
                    
            except Exception as e:
                logger.error(f"检索文档片段失败: {str(e)}")
                results.append({
                    'item_id': str(i + 1),
                    'chapter': chapter,
                    'name': name,
                    'required': required,
                    'item_type': item_type,
                    'ai_applicable': ai_applicable,
                    'description': description,
                    'completeness_status': '检查失败',
                    'completeness_score': 0.0,
                    'evidence': '',
                    'detailed_result': f'检索失败: {str(e)}'
                })
                
        except Exception as e:
            logger.error(f"检查项目 {i+1} 时发生错误: {str(e)}")
            results.append({
                'item_id': str(i + 1),
                'chapter': item.get('章节', ''),
                'name': item.get('名称', ''),
                'required': required,
                'item_type': item.get('类型', ''),
                'ai_applicable': item.get('AI适用', '否'),
                'description': item.get('说明', ''),
                'completeness_status': '检查失败',
                'completeness_score': 0.0,
                'evidence': '',
                'detailed_result': f'检查过程出错: {str(e)}'
            })
    
    return results

def perform_chapter_by_chapter_structure_check(toc_items, auditor, chat_model, top_k):
    """逐章节检查模式"""
    results = []
    
    # 按章节分组
    chapters = {}
    for i, item in enumerate(toc_items):
        chapter_num = item.get('章节', '')
        # 提取主章节号（如"1.1.2"的"1"）
        chapter_prefix = chapter_num.split('.')[0] if '.' in chapter_num else chapter_num
        if chapter_prefix not in chapters:
            chapters[chapter_prefix] = []
        chapters[chapter_prefix].append((i, item))
    
    # 逐章节检查
    for chapter_prefix, chapter_items in chapters.items():
        try:
            # 构建章节级检索查询
            chapter_query_parts = []
            for _, item in chapter_items:
                chapter_query_parts.append(f"{item.get('章节', '')} {item.get('名称', '')}")
            
            chapter_query = f"第{chapter_prefix}章 " + " ".join(chapter_query_parts[:5])  # 限制查询长度
            
            # 检索相关文档片段
            try:
                similar_chunks_results = auditor.search_similar_chunks(chapter_query, top_k=min(top_k * 3, 20))
                
                # 过滤相似度低的结果
                similar_chunks = []
                for result in similar_chunks_results:
                    if result['similarity'] > 0.1:
                        similar_chunks.append((result['text'], result['similarity']))
                
                # 构建章节证据文本
                evidence_text = ""
                if similar_chunks:
                    evidence_chunks = []
                    for chunk, similarity in similar_chunks:
                        evidence_chunks.append(f"相关度{similarity:.3f}: {chunk}")
                    evidence_text = '\n'.join(evidence_chunks[:10])  # 章节级分析用更多证据
                
                # 使用AI进行章节级批量分析
                chapter_analysis = analyze_chapter_structure_completeness_batch(
                    chapter_items, evidence_text, chat_model, auditor
                )
                
                # 将章节分析结果分配到各个项目
                for i, item in chapter_items:
                    item_result = extract_item_result_from_chapter_analysis(
                        i + 1, item, chapter_analysis, evidence_text
                    )
                    results.append(item_result)
                    
            except Exception as e:
                logger.error(f"检索章节 {chapter_prefix} 失败: {str(e)}")
                # 章节检查失败时，为该章节的所有项目添加失败结果
                for i, item in chapter_items:
                    results.append({
                        'item_id': str(i + 1),
                        'chapter': item.get('章节', ''),
                        'name': item.get('名称', ''),
                        'required': item.get('必有', '否'),
                        'item_type': item.get('类型', ''),
                        'ai_applicable': item.get('AI适用', '否'),
                        'description': item.get('说明', ''),
                        'completeness_status': '检查失败',
                        'completeness_score': 0.0,
                        'evidence': '',
                        'detailed_result': f'章节检索失败: {str(e)}'
                    })
                
        except Exception as e:
            logger.error(f"检查章节 {chapter_prefix} 时发生错误: {str(e)}")
            # 章节检查失败时，为该章节的所有项目添加失败结果
            for i, item in chapter_items:
                results.append({
                    'item_id': str(i + 1),
                    'chapter': item.get('章节', ''),
                    'name': item.get('名称', ''),
                    'required': item.get('必有', '否'),
                    'item_type': item.get('类型', ''),
                    'ai_applicable': item.get('AI适用', '否'),
                    'description': item.get('说明', ''),
                    'completeness_status': '检查失败',
                    'completeness_score': 0.0,
                    'evidence': '',
                    'detailed_result': f'章节检查过程出错: {str(e)}'
                })
    
    # 按原始顺序排序
    results.sort(key=lambda x: int(x['item_id']))
    return results

def analyze_structure_completeness_single(item, evidence_text, chat_model, auditor):
    """使用AI分析单个项目的结构完整性"""
    chapter = item.get('章节', '')
    name = item.get('名称', '')
    required = item.get('必有', '否')
    item_type = item.get('类型', '')
    description = item.get('说明', '')
    
    messages = [
        {
            "role": "system",
            "content": "你是一位专业的建筑工程施工方案审查专家。请分析文档中是否包含指定的目录项内容，并评估其完整性。"
        },
        {
            "role": "user",
            "content": f"""请分析文档中是否包含指定的目录项内容。

目录项信息：
- 章节：{chapter}
- 名称：{name}
- 是否必有：{required}
- 类型：{item_type}
- 检查说明：{description}

文档相关内容：
{evidence_text}

请按以下格式回答：
1. 完整性状态：[完整/部分完整/缺失]
2. 完整性评分：[0.0-1.0之间的数值]
3. 分析说明：[详细说明分析过程和判断依据]

评估标准：
- 完整：目录项内容齐全，满足施工方案要求
- 部分完整：目录项内容存在但不够完整或详细
- 缺失：未找到相关内容或内容严重不足
"""
        }
    ]
    
    try:
        llm_response = auditor.embedder.generate_text(
            messages=messages,
            model=chat_model,
            temperature=0.1
        )
        
        # 解析AI返回结果
        completeness_status = "缺失"  # 默认状态
        completeness_score = 0.0  # 默认评分
        
        # 简单的文本解析来提取完整性状态
        response_lower = llm_response.lower()
        if "完整" in response_lower and "部分" not in response_lower:
            completeness_status = "完整"
        elif "部分完整" in response_lower or "部分" in response_lower:
            completeness_status = "部分完整"
        elif "缺失" in response_lower:
            completeness_status = "缺失"
        
        # 尝试提取完整性评分
        import re
        patterns = [
            r'完整性评分[：:]\s*([0-9.]+)',
            r'评分[：:]\s*([0-9.]+)',
            r'2\.\s*完整性评分[：:]\s*([0-9.]+)'
        ]
        
        for pattern in patterns:
            score_match = re.search(pattern, llm_response)
            if score_match:
                try:
                    completeness_score = float(score_match.group(1))
                    completeness_score = max(0.0, min(1.0, completeness_score))  # 确保在0-1范围内
                    break
                except ValueError:
                    continue
        
        return {
            'chapter': chapter,
            'name': name,
            'required': required,
            'item_type': item_type,
            'ai_applicable': item.get('AI适用', '否'),
            'description': description,
            'completeness_status': completeness_status,
            'completeness_score': round(completeness_score, 3),
            'evidence': evidence_text[:500] + '...' if len(evidence_text) > 500 else evidence_text,
            'detailed_result': llm_response
        }
        
    except Exception as e:
        logger.error(f"AI分析过程出错: {str(e)}")
        return simple_structure_check_single(item, evidence_text, 0)

def simple_structure_check_single(item, evidence_text, item_id):
    """简单的结构完整性检查（关键词匹配）"""
    chapter = item.get('章节', '')
    name = item.get('名称', '')
    required = item.get('必有', '否')
    
    # 基于关键词匹配的简单检查
    keywords = [name]
    if chapter:
        keywords.append(chapter)
    
    found_keywords = []
    for keyword in keywords:
        if keyword and keyword in evidence_text:
            found_keywords.append(keyword)
    
    if found_keywords:
        if len(found_keywords) == len([k for k in keywords if k]):
            status = '完整'
            score = 0.8
        else:
            status = '部分完整'
            score = 0.5
    else:
        if required == '是':
            status = '缺失'
            score = 0.0
        else:
            status = '部分完整'  # 非必须项目给予部分完整
            score = 0.3
    
    return {
        'item_id': str(item_id),
        'chapter': chapter,
        'name': name,
        'required': required,
        'item_type': item.get('类型', ''),
        'ai_applicable': item.get('AI适用', '否'),
        'description': item.get('说明', ''),
        'completeness_status': status,
        'completeness_score': score,
        'evidence': evidence_text[:500] + '...' if len(evidence_text) > 500 else evidence_text,
        'detailed_result': f"基于关键词匹配的简单检查。找到关键词: {', '.join(found_keywords) if found_keywords else '无'}"
    }

def analyze_chapter_structure_completeness_batch(chapter_items, evidence_text, chat_model, auditor):
    """使用AI分析整个章节的结构完整性"""
    chapter_info = []
    for i, item in chapter_items:
        chapter_info.append({
            'item_id': i + 1,
            'chapter': item.get('章节', ''),
            'name': item.get('名称', ''),
            'required': item.get('必有', '否'),
            'type': item.get('类型', ''),
            'description': item.get('说明', '')
        })
    
    messages = [
        {
            "role": "system",
            "content": "你是一位专业的建筑工程施工方案审查专家。请分析文档中指定章节的结构完整性，逐项评估每个目录项。"
        },
        {
            "role": "user",
            "content": f"""请分析文档中指定章节的结构完整性。

需要检查的章节项目：
{json.dumps(chapter_info, ensure_ascii=False, indent=2)}

文档相关内容：
{evidence_text}

请逐项分析每个目录项的完整性状况，并按以下格式回答：

章节整体分析：[对整个章节的总体评价]

逐项分析：
项目1 (ID: X):
- 完整性状态：[完整/部分完整/缺失]
- 完整性评分：[0.0-1.0之间的数值]
- 分析说明：[具体分析]

项目2 (ID: Y):
...

评估标准：
- 完整：目录项内容齐全，满足施工方案要求
- 部分完整：目录项内容存在但不够完整或详细
- 缺失：未找到相关内容或内容严重不足
"""
        }
    ]
    
    try:
        llm_response = auditor.embedder.generate_text(
            messages=messages,
            model=chat_model,
            temperature=0.1
        )
        
        return {
            'chapter_analysis': llm_response,
            'raw_response': llm_response
        }
        
    except Exception as e:
        logger.error(f"章节AI分析过程出错: {str(e)}")
        return None

def extract_item_result_from_chapter_analysis(item_id, item, chapter_analysis, evidence_text):
    """从章节分析结果中提取单个项目的结果"""
    if chapter_analysis and 'raw_response' in chapter_analysis:
        response = chapter_analysis['raw_response']
        
        # 尝试解析逐项分析结果
        import re
        
        # 查找对应项目的分析
        pattern = rf'项目.*?\(?ID:?\s*{item_id}\)?:?\s*(.*?)(?=项目.*?\(?ID:?\s*\d+\)?:|$)'
        match = re.search(pattern, response, re.DOTALL | re.IGNORECASE)
        
        if match:
            item_analysis = match.group(1).strip()
            
            # 解析完整性状态
            status = "缺失"
            if "完整性状态" in item_analysis:
                if "完整" in item_analysis and "部分" not in item_analysis:
                    status = "完整"
                elif "部分完整" in item_analysis or "部分" in item_analysis:
                    status = "部分完整"
                elif "缺失" in item_analysis:
                    status = "缺失"
            
            # 解析评分
            score = 0.0
            score_patterns = [
                r'完整性评分[：:]\s*([0-9.]+)',
                r'评分[：:]\s*([0-9.]+)'
            ]
            
            for pattern in score_patterns:
                score_match = re.search(pattern, item_analysis)
                if score_match:
                    try:
                        score = float(score_match.group(1))
                        score = max(0.0, min(1.0, score))  # 确保在0-1范围内
                        break
                    except ValueError:
                        continue
            
            return {
                'item_id': str(item_id),
                'chapter': item.get('章节', ''),
                'name': item.get('名称', ''),
                'required': item.get('必有', '否'),
                'item_type': item.get('类型', ''),
                'ai_applicable': item.get('AI适用', '否'),
                'description': item.get('说明', ''),
                'completeness_status': status,
                'completeness_score': round(score, 3),
                'evidence': evidence_text[:500] + '...' if len(evidence_text) > 500 else evidence_text,
                'detailed_result': item_analysis
            }
    
    # 如果没有找到对应的分析结果，使用简单检查
    return simple_structure_check_single(item, evidence_text, item_id) 