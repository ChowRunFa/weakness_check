# -*- coding: utf-8 -*-
"""
异步引用检查API
"""
import os
import json
import logging
import threading
import requests
import re
from datetime import datetime
from flask import Blueprint, request, jsonify
from flasgger import swag_from
from werkzeug.utils import secure_filename
from docx import Document

# 导入现有的工具类和方法
from objs.PlanAuditor import PlanAuditor
from objs.FileManager import FileManager
from utils.prompts import CONSTRUCTION_EXPERT_SYSTEM

# 导入数据库模块
from db import AsyncTaskDAO, DocumentDAO, CiteCheckDAO

# 导入swagger配置
from utils.swagger_configs.async_cite_check_swagger import (
    async_cite_check_swagger,
    get_cite_check_task_status_swagger,
    get_cite_check_result_swagger
)

# 创建蓝图
api_cite_check_async = Blueprint('api_cite_check_async', __name__)

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 全局配置
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx', 'doc', 'txt', 'pdf', 'json', 'jsonl'}
CACHE_DIR = 'cache'
CALLBACK_URL = '/test/cite/callback'  # 本地测试回调接口地址
DEFAULT_CALLBACK_BASE_URL = 'http://127.0.0.1:5000'  # 默认本地回调基础URL

# 确保目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CACHE_DIR, exist_ok=True)

def generate_timestamp_folder():
    """生成基于当前时间的文件夹名"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]
    return timestamp

def ensure_upload_subfolder():
    """创建并返回带时间戳的上传子文件夹路径"""
    timestamp_folder = generate_timestamp_folder()
    upload_subfolder = os.path.join(UPLOAD_FOLDER, timestamp_folder)
    os.makedirs(upload_subfolder, exist_ok=True)
    return upload_subfolder, timestamp_folder

def allowed_file(filename):
    """检查文件类型是否被允许"""
    if not filename or '.' not in filename:
        return False
    try:
        ext = filename.rsplit('.', 1)[1].lower()
        return ext in ALLOWED_EXTENSIONS
    except (IndexError, AttributeError):
        return False

def extract_text_from_docx(file_path):
    """从docx文件提取文本，增加错误处理"""
    try:
        doc = Document(file_path)
        text_content = []
        
        # 提取段落文本
        try:
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
        except Exception as e:
            logger.warning(f"提取段落时出错: {e}")
        
        # 提取表格文本
        try:
            for table_idx, table in enumerate(doc.tables):
                try:
                    for row_idx, row in enumerate(table.rows):
                        try:
                            for cell_idx in range(len(table.columns)):
                                try:
                                    cell = row.cells[cell_idx]
                                    if cell.text and cell.text.strip():
                                        text_content.append(cell.text.strip())
                                except (IndexError, AttributeError):
                                    continue
                        except Exception as e:
                            logger.debug(f"跳过表格 {table_idx} 行 {row_idx}: {e}")
                            continue
                except Exception as e:
                    logger.debug(f"跳过表格 {table_idx}: {e}")
                    continue
        except Exception as e:
            logger.warning(f"提取表格时出错: {e}")
        
        result = "\n".join(text_content)
        
        if not result.strip():
            return "文档内容为空或无法提取文本"
        
        logger.info(f"成功提取文档内容，长度: {len(result)} 字符")
        return result
        
    except Exception as e:
        logger.error(f"提取docx文件文本时发生错误: {str(e)}")
        raise ValueError(f"无法提取docx文件内容: {str(e)}")

def send_callback(callback_url, task_id, status, data=None, error_message=None):
    """发送回调请求并更新数据库状态"""
    try:
        # 1. 先更新数据库状态
        try:
            if status == "success" and data:
                # 保存引用检查结果到数据库
                CiteCheckDAO.save_check_result(task_id, data)
                # 更新任务状态为成功
                AsyncTaskDAO.update_task_status(
                    task_id, 
                    status, 
                    result_data=data.get('summary') if data else None
                )
            else:
                # 更新任务状态
                AsyncTaskDAO.update_task_status(
                    task_id, 
                    status, 
                    error_message=error_message
                )
            logger.info(f"数据库状态更新成功，任务ID: {task_id}, 状态: {status}")
        except Exception as db_error:
            logger.error(f"更新数据库状态失败: {str(db_error)}, 任务ID: {task_id}")
        
        # 2. 准备回调数据
        callback_data = {
            "task_id": task_id,
            "status": status,  # "success", "failed", "processing"
            "timestamp": datetime.now().isoformat(),
            "data": data,
            "error_message": error_message
        }
        
        headers = {
            'Content-Type': 'application/json'
        }
        
        logger.info(f"发送回调到: {callback_url}, 任务ID: {task_id}, 状态: {status}")
        
        # 3. 发送POST请求到回调接口
        response = requests.post(
            callback_url, 
            json=callback_data, 
            headers=headers,
            timeout=30
        )
        
        if response.status_code == 200:
            logger.info(f"回调发送成功，任务ID: {task_id}")
        else:
            logger.warning(f"回调发送失败，状态码: {response.status_code}, 任务ID: {task_id}")
            
    except Exception as e:
        logger.error(f"发送回调时发生错误: {str(e)}, 任务ID: {task_id}")

def async_cite_check_worker(task_params):
    """异步执行引用检查的工作函数"""
    task_id = task_params['task_id']
    callback_url_full = task_params['callback_url']
    
    try:
        logger.info(f"开始异步执行引用检查，任务ID: {task_id}")
        
        # 发送处理中状态回调
        send_callback(callback_url_full, task_id, "processing", {"message": "开始执行引用检查"})
        
        # 执行引用检查
        result = perform_cite_check_internal(task_params)
        
        # 发送成功回调
        send_callback(callback_url_full, task_id, "success", result)
        
        logger.info(f"引用检查完成，任务ID: {task_id}")
        
    except Exception as e:
        error_msg = f"引用检查执行失败: {str(e)}"
        logger.error(f"{error_msg}, 任务ID: {task_id}")
        
        # 发送失败回调
        send_callback(callback_url_full, task_id, "failed", error_message=error_msg)

def perform_cite_check_internal(task_params):
    """内部执行引用检查的具体逻辑"""
    cite_list_path = task_params['cite_list_path']
    document_path = task_params['document_path']
    cite_list_filename = task_params['cite_list_filename']
    document_filename = task_params['document_filename']
    embedding_model = task_params['embedding_model']
    chat_model = task_params['chat_model']
    top_k = task_params['top_k']
    openai_api_key = task_params['openai_api_key']
    openai_api_base = task_params['openai_api_base']
    timestamp = task_params['timestamp']
    
    # 解析引用检查文件
    citation_items = []
    try:
        with open(cite_list_path, 'r', encoding='utf-8') as f:
            if cite_list_filename.endswith('.jsonl'):
                # JSONL格式，每行一个JSON对象
                for line in f:
                    line = line.strip()
                    if line:
                        citation_items.append(json.loads(line))
            else:
                # JSON格式，可能是数组或单个对象
                content = json.load(f)
                if isinstance(content, list):
                    citation_items = content
                else:
                    citation_items = [content]
    except Exception as e:
        raise Exception(f'解析引用检查文件失败: {str(e)}')
    
    if not citation_items:
        raise Exception('引用检查文件为空或格式不正确')
    
    # 提取文档内容
    try:
        document_content = extract_text_from_docx(document_path)
    except Exception as e:
        raise Exception(f'提取文档内容失败: {str(e)}')
    
    # 创建临时引用检查文件（JSONL格式）
    temp_cite_list_path = os.path.join(os.path.dirname(cite_list_path), 'temp_cite_list.jsonl')
    with open(temp_cite_list_path, 'w', encoding='utf-8') as f:
        for item in citation_items:
            f.write(json.dumps(item, ensure_ascii=False) + '\n')
    
    # 初始化审查器
    try:
        auditor = PlanAuditor(
            plan_content=document_content,
            check_list_file=temp_cite_list_path,
            embedding_model=embedding_model,
            openai_api_key=openai_api_key,
            openai_api_base=openai_api_base,
            cache_dir=CACHE_DIR,
            original_filename=document_filename
        )
        
        # 构建嵌入
        plan_id = auditor.build_or_load_embeddings()
        
    except Exception as e:
        raise Exception(f'初始化审查器失败: {str(e)}')
    
    # 逐个检查每个引用条目
    citation_results = []
    
    for i, citation_item in enumerate(citation_items, 1):
        try:
            logger.info(f"正在检查第 {i}/{len(citation_items)} 个引用条目: {citation_item}")
            
            # 获取引用条目信息（支持多种字段格式）
            # 学术文献格式
            citation_id = citation_item.get('id', citation_item.get('citation_id', str(i)))
            title = citation_item.get('title', citation_item.get('标题', ''))
            authors = citation_item.get('authors', citation_item.get('作者', ''))
            publication = citation_item.get('publication', citation_item.get('出版物', ''))
            year = citation_item.get('year', citation_item.get('年份', ''))
            citation_text = citation_item.get('citation_text', citation_item.get('引用文本', ''))
            
            # 标准规范格式（优先使用标准规范字段）
            standard_code = citation_item.get('标准编号', '')
            standard_name = citation_item.get('标准名称', '')
            issuing_dept = citation_item.get('发布部门', '')
            implementation_date = citation_item.get('实施日期', '')
            status = citation_item.get('状态', '')
            
            # 如果是标准规范格式，重新组织字段
            if standard_code or standard_name:
                citation_id = citation_id or standard_code
                title = title or standard_name
                authors = authors or issuing_dept
                publication = publication or '标准规范'
                if implementation_date and not year:
                    # 从实施日期提取年份
                    try:
                        year = implementation_date.split('-')[0]
                    except:
                        year = ''
            
            # 如果没有明确的引用文本，则从其他字段构建查询
            if not citation_text:
                search_components = []
                if standard_code:
                    search_components.append(standard_code)
                if standard_name:
                    search_components.append(standard_name)
                elif title:
                    search_components.append(title)
                if issuing_dept:
                    search_components.append(issuing_dept)
                elif authors:
                    search_components.append(authors)
                if implementation_date:
                    search_components.append(implementation_date)
                elif year:
                    search_components.append(str(year))
                
                citation_text = ' '.join(search_components)
            
            if not citation_text.strip():
                citation_results.append({
                    'citation_id': citation_id,
                    'title': title,
                    'authors': authors,
                    'publication': publication,
                    'year': year,
                    'standard_code': standard_code,
                    'standard_name': standard_name,
                    'issuing_dept': issuing_dept,
                    'implementation_date': implementation_date,
                    'status': status,
                    'citation_text': citation_text,
                    'evidence': '',
                    'citation_status': '检查失败',
                    'accuracy_score': 0.0,
                    'detailed_result': '引用条目信息不完整，无法进行检查',
                    'chunk_count': 0,
                    'error': '缺少可检索的引用文本信息'
                })
                continue
            
            # 搜索相关文本片段
            similar_chunks = auditor.search_similar_chunks(citation_text, top_k=top_k)
            
            # 组合检索到的文本作为证据
            evidence_texts = []
            for chunk_dict in similar_chunks:
                chunk_text = chunk_dict.get("text", "")
                similarity = chunk_dict.get("similarity", 0.0)
                evidence_texts.append(f"相关度{similarity:.3f}: {chunk_text}")
            
            evidence = "\n".join(evidence_texts)
            
            # 使用大模型进行引用检查
            if similar_chunks:
                # 构建上下文内容
                context = "\n".join([chunk_dict.get("text", "") for chunk_dict in similar_chunks])
                
                # 构建引用检查prompt（针对标准规范或学术文献）
                if standard_code or standard_name:
                    # 标准规范检查prompt
                    messages = [
                        {
                            "role": "system",
                            "content": "你是一位专业的工程技术文档审查专家。请基于提供的文档内容，判断文档是否正确引用了指定的技术标准规范。"
                        },
                        {
                            "role": "user",
                            "content": f"""
请分析以下技术文档内容，判断是否正确引用了指定的技术标准规范。

【待检查的标准规范】:
- 标准编号: {standard_code}
- 标准名称: {standard_name}
- 发布部门: {issuing_dept}
- 实施日期: {implementation_date}
- 状态: {status}
- 查询文本: {citation_text}

【文档相关内容】:
{context}

请按以下格式回答：
1. 引用状态：[正确引用/缺失引用/引用有误/引用不完整]
2. 准确性评分：[0.0-1.0之间的数值，表示引用的准确性]
3. 分析说明：[详细说明分析过程和判断依据]

判断标准：
- 正确引用：文档中明确提及了该标准的编号、名称等关键信息
- 缺失引用：文档中应该引用该标准但完全没有提及
- 引用有误：文档中提及了该标准但信息有错误（编号错误、名称错误等）
- 引用不完整：文档中提及了该标准但引用信息不完整（如只有编号没有名称）
- 准确性评分反映标准引用的完整性和准确性
"""
                        }
                    ]
                else:
                    # 学术文献检查prompt
                    messages = [
                        {
                            "role": "system",
                            "content": "你是一位专业的学术引用检查专家。请基于提供的文档内容，判断文档是否正确引用了指定的学术条目。"
                        },
                        {
                            "role": "user",
                            "content": f"""
请分析以下文档内容，判断是否正确引用了指定的学术条目。

【待检查的引用条目】:
- 标题: {title}
- 作者: {authors}
- 出版物: {publication}
- 年份: {year}
- 查询文本: {citation_text}

【文档相关内容】:
{context}

请按以下格式回答：
1. 引用状态：[正确引用/缺失引用/引用有误/引用不完整]
2. 准确性评分：[0.0-1.0之间的数值，表示引用的准确性]
3. 分析说明：[详细说明分析过程和判断依据]

判断标准：
- 正确引用：文档中包含了该条目的准确引用信息（作者、标题、年份等）
- 缺失引用：文档中应该引用但完全没有提及该条目
- 引用有误：文档中提及了该条目但信息有错误（作者名错误、年份错误等）
- 引用不完整：文档中提及了该条目但引用信息不完整
- 准确性评分反映引用信息的完整性和准确性
"""
                        }
                    ]
                
                # 调用大模型进行判断
                try:
                    llm_response = auditor.embedder.generate_text(
                        messages=messages,
                        model=chat_model,
                        temperature=0.1
                    )
                    
                    # 解析大模型回复
                    citation_status = "缺失引用"  # 默认状态
                    accuracy_score = 0.0  # 默认评分
                    
                    # 简单的文本解析来提取引用状态
                    response_lower = llm_response.lower()
                    if "正确引用" in response_lower:
                        citation_status = "正确引用"
                    elif "引用有误" in response_lower:
                        citation_status = "引用有误"
                    elif "引用不完整" in response_lower:
                        citation_status = "引用不完整"
                    elif "缺失引用" in response_lower:
                        citation_status = "缺失引用"
                    
                    # 尝试提取准确性评分
                    patterns = [
                        r'准确性评分[：:]\s*([0-9.]+)',
                        r'评分[：:]\s*([0-9.]+)',
                        r'2\.\s*准确性评分[：:]\s*([0-9.]+)'
                    ]
                    
                    for pattern in patterns:
                        score_match = re.search(pattern, llm_response)
                        if score_match:
                            try:
                                accuracy_score = float(score_match.group(1))
                                accuracy_score = max(0.0, min(1.0, accuracy_score))  # 确保在0-1范围内
                                break
                            except ValueError:
                                continue
                    
                    detailed_result = llm_response
                    
                except Exception as e:
                    logger.error(f"大模型调用失败: {str(e)}")
                    # 降级到简单判断
                    citation_status = "检查失败"
                    accuracy_score = 0.0
                    detailed_result = f"大模型调用失败，无法完成引用检查: {str(e)}"
                    
                citation_results.append({
                    'citation_id': citation_id,
                    'title': title,
                    'authors': authors,
                    'publication': publication,
                    'year': year,
                    'standard_code': standard_code,
                    'standard_name': standard_name,
                    'issuing_dept': issuing_dept,
                    'implementation_date': implementation_date,
                    'status': status,
                    'citation_text': citation_text,
                    'evidence': evidence,
                    'citation_status': citation_status,
                    'accuracy_score': round(accuracy_score, 3),
                    'detailed_result': detailed_result,
                    'chunk_count': len(similar_chunks)
                })
            
        except Exception as e:
            logger.error(f"检查第 {i} 个引用条目时发生错误: {str(e)}")
            citation_results.append({
                'citation_id': citation_item.get('id', citation_item.get('citation_id', citation_item.get('标准编号', str(i)))),
                'title': citation_item.get('title', citation_item.get('标题', citation_item.get('标准名称', ''))),
                'authors': citation_item.get('authors', citation_item.get('作者', citation_item.get('发布部门', ''))),
                'publication': citation_item.get('publication', citation_item.get('出版物', '标准规范')),
                'year': citation_item.get('year', citation_item.get('年份', '')),
                'standard_code': citation_item.get('标准编号', ''),
                'standard_name': citation_item.get('标准名称', ''),
                'issuing_dept': citation_item.get('发布部门', ''),
                'implementation_date': citation_item.get('实施日期', ''),
                'status': citation_item.get('状态', ''),
                'citation_text': citation_item.get('citation_text', citation_item.get('引用文本', '')),
                'evidence': '',
                'citation_status': '检查失败',
                'accuracy_score': 0.0,
                'detailed_result': f'检查过程中发生错误: {str(e)}',
                'chunk_count': 0,
                'error': str(e)
            })
    
    # 计算总体统计
    total_citations = len(citation_results)
    properly_cited = len([r for r in citation_results if r['citation_status'] == '正确引用'])
    missing_citations = len([r for r in citation_results if r['citation_status'] == '缺失引用'])
    incorrectly_cited = len([r for r in citation_results if r['citation_status'] in ['引用有误', '引用不完整']])
    failed_checks = len([r for r in citation_results if r['citation_status'] == '检查失败'])
    
    return {
        'summary': {
            'total_citations': total_citations,
            'properly_cited': properly_cited,
            'missing_citations': missing_citations,
            'incorrectly_cited': incorrectly_cited,
            'failed_checks': failed_checks,
            'citation_rate': round(properly_cited / total_citations * 100, 2) if total_citations > 0 else 0
        },
        'citation_results': citation_results,
        'document_filename': document_filename,
        'cite_list_filename': cite_list_filename,
        'plan_id': plan_id,
        'upload_folder': timestamp
    }

@api_cite_check_async.route('/async_cite_check', methods=['POST'])
@swag_from(async_cite_check_swagger)
def async_cite_check():
    """
    异步引用检查接口
    
    立即返回调用结果，后台异步执行检查，完成后回调指定接口
    """
    try:
        # 检查必要参数
        scheme_id = request.form.get('schemeId')
        file_path = request.form.get('filePath')
        scheme_name = request.form.get('schemeName')
        cite_list_path = request.form.get('citeListPath')
        
        # 验证必要参数
        if not scheme_id:
            return jsonify({
                'code': 400,
                'message': '方案ID不能为空',
                'data': {'result': 'false'}
            }), 400
        
        if not file_path:
            return jsonify({
                'code': 400,
                'message': '文档文件路径不能为空',
                'data': {'result': 'false'}
            }), 400
        
        if not scheme_name:
            return jsonify({
                'code': 400,
                'message': '方案名称不能为空',
                'data': {'result': 'false'}
            }), 400
        
        if not cite_list_path:
            return jsonify({
                'code': 400,
                'message': '引用检查文件路径不能为空',
                'data': {'result': 'false'}
            }), 400
        
        # 验证引用检查文件是否存在
        if not os.path.exists(cite_list_path):
            return jsonify({
                'code': 400,
                'message': f'引用检查文件不存在: {cite_list_path}',
                'data': {'result': 'false'}
            }), 400
        
        # 验证引用检查文件类型
        if not cite_list_path.lower().endswith(('.json', '.jsonl')):
            return jsonify({
                'code': 400,
                'message': '引用检查文件必须是JSON或JSONL格式',
                'data': {'result': 'false'}
            }), 400
        
        # 验证待检查文档是否存在
        if not os.path.exists(file_path):
            return jsonify({
                'code': 400,
                'message': f'文档文件不存在: {file_path}',
                'data': {'result': 'false'}
            }), 400
        
        # 验证文档文件类型
        if not file_path.lower().endswith('.docx'):
            return jsonify({
                'code': 400,
                'message': '文档必须是DOCX格式',
                'data': {'result': 'false'}
            }), 400
        
        # 创建临时文件夹用于任务标识
        upload_subfolder, timestamp_folder = ensure_upload_subfolder()
        
        # 直接使用提供的文件路径
        document_path = file_path
        document_filename = os.path.basename(file_path)
        cite_list_filename = os.path.basename(cite_list_path)
        
        # 获取配置参数
        callback_base_url = request.form.get('callback_base_url', DEFAULT_CALLBACK_BASE_URL)
        embedding_model = request.form.get('embedding_model', 'nomic-embed-text:latest')
        chat_model = request.form.get('chat_model', 'qwen2.5:32b')
        top_k = int(request.form.get('top_k', 5))
        openai_api_key = request.form.get('openai_api_key', 'ollama')
        openai_api_base = request.form.get('openai_api_base', 'http://59.77.7.24:11434/v1/')
        
        # 生成任务ID
        timestamp = generate_timestamp_folder()
        task_id = f"cite_check_{timestamp}"
        
        # 构建回调URL
        callback_url_full = f"{callback_base_url.rstrip('/')}{CALLBACK_URL}"
        
        # 保存任务到数据库
        try:
            # 创建异步任务记录
            request_params = {
                'scheme_id': scheme_id,
                'file_path': file_path,
                'scheme_name': scheme_name,
                'cite_list_path': cite_list_path,
                'cite_list_filename': cite_list_filename,
                'document_filename': document_filename,
                'embedding_model': embedding_model,
                'chat_model': chat_model,
                'top_k': top_k,
                'callback_base_url': callback_base_url
            }
            
            task = AsyncTaskDAO.create_task(
                task_id=task_id,
                task_type='cite_check',
                callback_url=callback_url_full,
                request_params=request_params,
                scheme_id=int(scheme_id),
                scheme_name=scheme_name
            )
            
            if not task:
                return jsonify({
                    'code': 500,
                    'message': '创建任务记录失败',
                    'data': {'result': 'false'}
                }), 500
            
            # 保存文档信息
            cite_list_file_size = os.path.getsize(cite_list_path) if os.path.exists(cite_list_path) else 0
            document_file_size = os.path.getsize(document_path) if os.path.exists(document_path) else 0
            
            # 保存引用检查文件引用（引用已存在的引用检查文件）
            DocumentDAO.save_document_reference(
                task_id=task_id,
                scheme_id=int(scheme_id),
                original_filename=cite_list_filename,
                saved_filename=cite_list_filename,
                file_path=cite_list_path,
                file_size=cite_list_file_size,
                file_type='cite_list',
                reference_folder=f'scheme_{scheme_id}'
            )
            
            # 保存文档文件引用（引用已存在的文档）
            DocumentDAO.save_document_reference(
                task_id=task_id,
                scheme_id=int(scheme_id),
                original_filename=document_filename,
                saved_filename=document_filename,
                file_path=document_path,
                file_size=document_file_size,
                file_type='document',
                reference_folder=f'scheme_{scheme_id}'
            )
            
            logger.info(f"任务和文档信息已保存到数据库，任务ID: {task_id}")
            
        except Exception as db_error:
            logger.error(f"保存任务到数据库失败: {str(db_error)}")
            return jsonify({
                'code': 500,
                'message': f'保存任务失败: {str(db_error)}',
                'data': {'result': 'false'}
            }), 500
        
        # 准备异步任务参数
        task_params = {
            'task_id': task_id,
            'scheme_id': scheme_id,
            'scheme_name': scheme_name,
            'callback_url': callback_url_full,
            'cite_list_path': cite_list_path,
            'document_path': document_path,
            'cite_list_filename': cite_list_filename,
            'document_filename': document_filename,
            'embedding_model': embedding_model,
            'chat_model': chat_model,
            'top_k': top_k,
            'openai_api_key': openai_api_key,
            'openai_api_base': openai_api_base,
            'timestamp': timestamp_folder
        }
        
        # 启动异步任务
        thread = threading.Thread(
            target=async_cite_check_worker,
            args=(task_params,),
            daemon=True
        )
        thread.start()
        
        logger.info(f"异步引用检查任务已启动，任务ID: {task_id}")
        
        # 立即返回成功响应
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': {
                'result': 'true',
                'task_id': task_id,
                'callback_url': callback_url_full,
                'estimated_time': '预计3-10分钟内完成'
            }
        }), 200
        
    except Exception as e:
        logger.error(f"异步引用检查API发生错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器内部错误: {str(e)}',
            'data': {'result': 'false'}
        }), 500

# 任务状态查询API
@api_cite_check_async.route('/async_cite_check/status/<task_id>', methods=['GET'])
@swag_from(get_cite_check_task_status_swagger)
def get_task_status(task_id):
    """查询异步任务状态"""
    try:
        task = AsyncTaskDAO.get_task_by_id(task_id)
        if not task:
            return jsonify({
                'code': 404,
                'message': '任务不存在',
                'data': None
            }), 404
        
        task_data = task.to_dict()
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': task_data
        }), 200
        
    except Exception as e:
        logger.error(f"查询任务状态失败: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'查询失败: {str(e)}',
            'data': None
        }), 500

# 检查结果查询API
@api_cite_check_async.route('/async_cite_check/result/<task_id>', methods=['GET'])
@swag_from(get_cite_check_result_swagger)
def get_check_result(task_id):
    """获取引用检查结果"""
    try:
        # 使用CiteCheckDAO获取结果
        result = CiteCheckDAO.get_check_result(task_id)
        if not result:
            return jsonify({
                'code': 404,
                'message': '检查结果不存在或任务未完成',
                'data': None
            }), 404
        
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': result
        }), 200
        
    except Exception as e:
        logger.error(f"获取检查结果失败: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'获取结果失败: {str(e)}',
            'data': None
        }), 500

# 本地测试回调接口
@api_cite_check_async.route('/test/cite/callback', methods=['POST'])
def test_callback():
    """本地测试回调接口"""
    try:
        # 获取回调数据
        callback_data = request.get_json()
        
        if not callback_data:
            logger.warning("回调接收失败：未收到有效的JSON数据")
            return jsonify({
                'code': 400,
                'message': '无效的回调数据',
                'data': None
            }), 400
        
        task_id = callback_data.get('task_id')
        status = callback_data.get('status')
        timestamp = callback_data.get('timestamp')
        data = callback_data.get('data')
        error_message = callback_data.get('error_message')
        
        logger.info(f"收到引用检查回调请求 - 任务ID: {task_id}, 状态: {status}, 时间: {timestamp}")
        
        # 打印回调数据详情
        if status == "success" and data:
            summary = data.get('summary', {})
            logger.info(f"任务 {task_id} 执行成功:")
            logger.info(f"  - 总引用数: {summary.get('total_citations', 0)}")
            logger.info(f"  - 正确引用数: {summary.get('properly_cited', 0)}")
            logger.info(f"  - 缺失引用数: {summary.get('missing_citations', 0)}")
            logger.info(f"  - 错误引用数: {summary.get('incorrectly_cited', 0)}")
            logger.info(f"  - 检查失败数: {summary.get('failed_checks', 0)}")
            logger.info(f"  - 引用正确率: {summary.get('citation_rate', 0)}%")
        elif status == "failed":
            logger.error(f"任务 {task_id} 执行失败: {error_message}")
        elif status == "processing":
            logger.info(f"任务 {task_id} 正在处理中...")
        
        # 模拟处理回调数据
        response_data = {
            'received_task_id': task_id,
            'received_status': status,
            'received_at': datetime.now().isoformat(),
            'callback_processed': True
        }
        
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': response_data
        }), 200
        
    except Exception as e:
        logger.error(f"处理回调请求时发生错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'处理回调失败: {str(e)}',
            'data': None
        }), 500
