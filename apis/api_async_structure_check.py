# -*- coding: utf-8 -*-
"""
异步文档结构完整性检查API
"""
import os
import json
import logging
import threading
import requests
from datetime import datetime
from flask import Blueprint, request, jsonify
from flasgger import swag_from
from werkzeug.utils import secure_filename
from docx import Document

# 导入现有的工具类和方法
from objs.PlanAuditor import PlanAuditor
from objs.FileManager import FileManager
from utils.prompts import CONSTRUCTION_EXPERT_SYSTEM

# 导入数据库模块
from db import AsyncTaskDAO, StructureCheckDAO, DocumentDAO

# 导入swagger配置
from utils.swagger_configs.async_structure_check_swagger import (
    async_structure_check_swagger,
    get_task_status_swagger,
    get_check_result_swagger,
    list_tasks_swagger
)

# 创建蓝图
api_async_structure_check = Blueprint('api_async_structure_check', __name__)

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 全局配置
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx', 'doc', 'txt', 'pdf', 'json', 'jsonl'}
CACHE_DIR = 'cache'
CALLBACK_URL = '/test/callback'  # 本地测试回调接口地址
DEFAULT_CALLBACK_BASE_URL = 'http://127.0.0.1:5000'  # 默认本地回调基础URL

# 确保目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CACHE_DIR, exist_ok=True)

def generate_timestamp_folder():
    """生成基于当前时间的文件夹名"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]
    return timestamp

def ensure_upload_subfolder():
    """创建并返回带时间戳的上传子文件夹路径"""
    timestamp_folder = generate_timestamp_folder()
    upload_subfolder = os.path.join(UPLOAD_FOLDER, timestamp_folder)
    os.makedirs(upload_subfolder, exist_ok=True)
    return upload_subfolder, timestamp_folder

def allowed_file(filename):
    """检查文件类型是否被允许"""
    if not filename or '.' not in filename:
        return False
    try:
        ext = filename.rsplit('.', 1)[1].lower()
        return ext in ALLOWED_EXTENSIONS
    except (IndexError, AttributeError):
        return False

def extract_text_from_docx(file_path):
    """从docx文件提取文本，增加错误处理"""
    try:
        doc = Document(file_path)
        text_content = []
        
        # 提取段落文本
        try:
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
        except Exception as e:
            logger.warning(f"提取段落时出错: {e}")
        
        # 提取表格文本
        try:
            for table_idx, table in enumerate(doc.tables):
                try:
                    for row_idx, row in enumerate(table.rows):
                        try:
                            for cell_idx in range(len(table.columns)):
                                try:
                                    cell = row.cells[cell_idx]
                                    if cell.text and cell.text.strip():
                                        text_content.append(cell.text.strip())
                                except (IndexError, AttributeError):
                                    continue
                        except Exception as e:
                            logger.debug(f"跳过表格 {table_idx} 行 {row_idx}: {e}")
                            continue
                except Exception as e:
                    logger.debug(f"跳过表格 {table_idx}: {e}")
                    continue
        except Exception as e:
            logger.warning(f"提取表格时出错: {e}")
        
        result = "\n".join(text_content)
        
        if not result.strip():
            return "文档内容为空或无法提取文本"
        
        logger.info(f"成功提取文档内容，长度: {len(result)} 字符")
        return result
        
    except Exception as e:
        logger.error(f"提取docx文件文本时发生错误: {str(e)}")
        raise ValueError(f"无法提取docx文件内容: {str(e)}")

def send_callback(callback_url, task_id, status, data=None, error_message=None):
    """发送回调请求并更新数据库状态"""
    try:
        # 1. 先更新数据库状态
        try:
            if status == "success" and data:
                # 保存结构检查结果到数据库
                StructureCheckDAO.save_check_result(task_id, data)
                # 更新任务状态为成功
                AsyncTaskDAO.update_task_status(
                    task_id, 
                    status, 
                    result_data=data.get('summary') if data else None
                )
            else:
                # 更新任务状态
                AsyncTaskDAO.update_task_status(
                    task_id, 
                    status, 
                    error_message=error_message
                )
            logger.info(f"数据库状态更新成功，任务ID: {task_id}, 状态: {status}")
        except Exception as db_error:
            logger.error(f"更新数据库状态失败: {str(db_error)}, 任务ID: {task_id}")
        
        # 2. 准备回调数据
        callback_data = {
            "task_id": task_id,
            "status": status,  # "success", "failed", "processing"
            "timestamp": datetime.now().isoformat(),
            "data": data,
            "error_message": error_message
        }
        
        headers = {
            'Content-Type': 'application/json'
        }
        
        logger.info(f"发送回调到: {callback_url}, 任务ID: {task_id}, 状态: {status}")
        
        # 3. 发送POST请求到回调接口
        response = requests.post(
            callback_url, 
            json=callback_data, 
            headers=headers,
            timeout=30
        )
        
        if response.status_code == 200:
            logger.info(f"回调发送成功，任务ID: {task_id}")
        else:
            logger.warning(f"回调发送失败，状态码: {response.status_code}, 任务ID: {task_id}")
            
    except Exception as e:
        logger.error(f"发送回调时发生错误: {str(e)}, 任务ID: {task_id}")

def async_structure_check_worker(task_params):
    """异步执行结构完整性检查的工作函数"""
    task_id = task_params['task_id']
    callback_url_full = task_params['callback_url']
    
    try:
        logger.info(f"开始异步执行结构检查，任务ID: {task_id}")
        
        # 发送处理中状态回调
        send_callback(callback_url_full, task_id, "processing", {"message": "开始执行结构完整性检查"})
        
        # 执行结构完整性检查
        result = perform_structure_check_internal(task_params)
        
        # 发送成功回调
        send_callback(callback_url_full, task_id, "success", result)
        
        logger.info(f"结构检查完成，任务ID: {task_id}")
        
    except Exception as e:
        error_msg = f"结构检查执行失败: {str(e)}"
        logger.error(f"{error_msg}, 任务ID: {task_id}")
        
        # 发送失败回调
        send_callback(callback_url_full, task_id, "failed", error_message=error_msg)

def perform_structure_check_internal(task_params):
    """内部执行结构完整性检查的具体逻辑"""
    toc_list_path = task_params['toc_list_path']
    document_path = task_params['document_path']
    toc_list_filename = task_params['toc_list_filename']
    document_filename = task_params['document_filename']
    check_mode = task_params['check_mode']
    embedding_model = task_params['embedding_model']
    chat_model = task_params['chat_model']
    top_k = task_params['top_k']
    openai_api_key = task_params['openai_api_key']
    openai_api_base = task_params['openai_api_base']
    scheme_id = task_params['scheme_id']
    timestamp = task_params['timestamp']
    
    # 解析目录结构清单
    toc_items = []
    try:
        with open(toc_list_path, 'r', encoding='utf-8') as f:
            if toc_list_filename.lower().endswith('.json'):
                toc_items = json.load(f)
            else:  # jsonl
                for line in f:
                    line = line.strip()
                    if line:
                        toc_items.append(json.loads(line))
    except Exception as e:
        raise Exception(f'解析目录结构清单失败: {str(e)}')
    
    if not toc_items:
        raise Exception('目录结构清单为空')
    
    # 处理文档并生成嵌入向量
    try:
        # 提取文档文本
        doc_text = extract_text_from_docx(document_path)
        if not doc_text:
            raise Exception('无法从DOCX文档中提取文本内容')
        
        # 创建临时检查清单文件
        temp_dir = os.path.join(CACHE_DIR, f'scheme_{scheme_id}')
        os.makedirs(temp_dir, exist_ok=True)
        temp_check_list = os.path.join(temp_dir, 'temp_check_list.jsonl')
        with open(temp_check_list, 'w', encoding='utf-8') as f:
            f.write('{"temp": "temp"}\n')
        
        # 创建文档处理器
        auditor = PlanAuditor(
            plan_content=doc_text,
            check_list_file=temp_check_list,
            embedding_model=embedding_model,
            openai_api_key=openai_api_key,
            openai_api_base=openai_api_base,
            cache_dir=CACHE_DIR,
            original_filename=document_filename
        )
        
        # 清理可能的缓存冲突，强制重新构建嵌入向量
        import shutil
        import hashlib
        
        # 生成当前配置的唯一标识
        config_hash = hashlib.md5(f"{embedding_model}_{document_filename}_{len(doc_text)}".encode()).hexdigest()[:8]
        scheme_cache_dir = os.path.join(CACHE_DIR, f'scheme_{scheme_id}_{config_hash}')
        
        # 如果存在旧缓存且配置不匹配，清理缓存
        if os.path.exists(scheme_cache_dir):
            try:
                # 检查缓存是否与当前配置匹配
                metadata_file = os.path.join(scheme_cache_dir, 'metadata.json')
                if os.path.exists(metadata_file):
                    with open(metadata_file, 'r', encoding='utf-8') as f:
                        metadata = json.load(f)
                    if metadata.get('embedding_model') != embedding_model:
                        logger.info(f"嵌入模型已变更，清理缓存: {scheme_cache_dir}")
                        shutil.rmtree(scheme_cache_dir)
            except Exception as cache_error:
                logger.warning(f"检查缓存时出错，清理缓存: {str(cache_error)}")
                shutil.rmtree(scheme_cache_dir, ignore_errors=True)
        
        # 更新auditor使用新的缓存目录
        auditor.cache_dir = scheme_cache_dir
        os.makedirs(scheme_cache_dir, exist_ok=True)
        
        # 构建嵌入向量
        logger.info(f"开始构建嵌入向量，缓存目录: {scheme_cache_dir}")
        plan_id = auditor.build_or_load_embeddings()
        logger.info(f"嵌入向量构建完成，plan_id: {plan_id}")
        
    except Exception as e:
        raise Exception(f'文档处理失败: {str(e)}')
    
    # 根据检查模式进行结构完整性检查
    if check_mode == 'item_by_item':
        check_results = perform_item_by_item_structure_check(toc_items, auditor, chat_model, top_k)
    else:  # chapter_by_chapter
        check_results = perform_chapter_by_chapter_structure_check(toc_items, auditor, chat_model, top_k)
    
    # 计算统计信息
    total_items = len(check_results)
    complete_items = len([r for r in check_results if r['completeness_status'] == '完整'])
    missing_items = len([r for r in check_results if r['completeness_status'] == '缺失'])
    partial_items = len([r for r in check_results if r['completeness_status'] == '部分完整'])
    failed_checks = len([r for r in check_results if r['completeness_status'] == '检查失败'])
    
    completeness_rate = (complete_items / total_items * 100) if total_items > 0 else 0
    
    return {
        'summary': {
            'total_items': total_items,
            'complete_items': complete_items,
            'missing_items': missing_items,
            'partial_items': partial_items,
            'failed_checks': failed_checks,
            'completeness_rate': round(completeness_rate, 2)
        },
        'check_results': check_results,
        'check_mode': check_mode,
        'toc_list_filename': toc_list_filename,
        'document_filename': document_filename,
        'plan_id': plan_id,
        'scheme_id': scheme_id,
        'upload_folder': f'scheme_{scheme_id}'
    }

@api_async_structure_check.route('/async_structure_check', methods=['POST'])
@swag_from(async_structure_check_swagger)
def async_structure_check():
    """
    异步文档结构完整性检查接口
    
    立即返回调用结果，后台异步执行检查，完成后回调指定接口
    """
    try:
        # 获取请求参数
        scheme_id = request.form.get('schemeId')
        file_path = request.form.get('filePath')
        scheme_name = request.form.get('schemeName')
        file_url = request.form.get('fileUrl')
        
        # 验证必要参数
        if not scheme_id:
            return jsonify({
                'code': 400,
                'message': '方案ID不能为空',
                'data': {'result': 'false'}
            }), 400
        
        if not file_path:
            return jsonify({
                'code': 400,
                'message': '文档文件路径不能为空',
                'data': {'result': 'false'}
            }), 400
        
        if not scheme_name:
            return jsonify({
                'code': 400,
                'message': '方案名称不能为空',
                'data': {'result': 'false'}
            }), 400
        
        if not file_url:
            return jsonify({
                'code': 400,
                'message': '模板文件路径不能为空',
                'data': {'result': 'false'}
            }), 400
        
        # 验证文件是否存在
        if not os.path.exists(file_path):
            return jsonify({
                'code': 400,
                'message': f'文档文件不存在: {file_path}',
                'data': {'result': 'false'}
            }), 400
        
        if not os.path.exists(file_url):
            return jsonify({
                'code': 400,
                'message': f'模板文件不存在: {file_url}',
                'data': {'result': 'false'}
            }), 400
        
        # 验证文件类型
        if not file_path.lower().endswith('.docx'):
            return jsonify({
                'code': 400,
                'message': '文档必须是DOCX格式',
                'data': {'result': 'false'}
            }), 400
        
        if not file_url.lower().endswith(('.json', '.jsonl')):
            return jsonify({
                'code': 400,
                'message': '模板文件必须是JSON或JSONL格式',
                'data': {'result': 'false'}
            }), 400
        
        # 获取参数
        check_mode = request.form.get('check_mode', 'chapter_by_chapter')
        callback_base_url = request.form.get('callback_base_url', DEFAULT_CALLBACK_BASE_URL)
        embedding_model = request.form.get('embedding_model', 'nomic-embed-text:latest')
        chat_model = request.form.get('chat_model', 'qwen2.5:32b')
        top_k = int(request.form.get('top_k', 5))
        openai_api_key = request.form.get('openai_api_key', 'ollama')
        openai_api_base = request.form.get('openai_api_base', 'http://59.77.7.24:11434/v1/')
        
        # 验证检查模式
        if check_mode not in ['item_by_item', 'chapter_by_chapter']:
            return jsonify({
                'code': 400,
                'message': '检查模式必须是 item_by_item 或 chapter_by_chapter',
                'data': {'result': 'false'}
            }), 400
        
        # 生成任务ID
        timestamp = generate_timestamp_folder()
        task_id = f"{scheme_id}_{timestamp}"  # 使用方案ID+时间戳作为任务ID
        
        # 使用传递的文件路径
        toc_list_path = file_url  # 模板文件路径
        document_path = file_path  # 文档文件路径
        
        # 从路径中提取文件名
        toc_list_filename = os.path.basename(toc_list_path)
        document_filename = os.path.basename(document_path)
        
        # 构建回调URL
        callback_url_full = f"{callback_base_url.rstrip('/')}{CALLBACK_URL}"
        
        # 保存任务到数据库
        try:
            # 1. 创建异步任务记录
            request_params = {
                'scheme_id': scheme_id,
                'scheme_name': scheme_name,
                'file_path': file_path,
                'file_url': file_url,
                'check_mode': check_mode,
                'embedding_model': embedding_model,
                'chat_model': chat_model,
                'top_k': top_k,
                'toc_list_filename': toc_list_filename,
                'document_filename': document_filename,
                'callback_base_url': callback_base_url
            }
            
            task = AsyncTaskDAO.create_task(
                task_id=task_id,
                task_type='structure_check',
                callback_url=callback_url_full,
                request_params=request_params,
                scheme_id=int(scheme_id),
                scheme_name=scheme_name
            )
            
            if not task:
                return jsonify({
                    'code': 500,
                    'message': '创建任务记录失败',
                    'data': {'result': 'false'}
                }), 500
            
            # 2. 保存文档信息
            toc_file_size = os.path.getsize(toc_list_path) if os.path.exists(toc_list_path) else 0
            doc_file_size = os.path.getsize(document_path) if os.path.exists(document_path) else 0
            
            # 保存目录清单文件引用
            DocumentDAO.save_document_reference(
                task_id=task_id,
                scheme_id=int(scheme_id),
                original_filename=toc_list_filename,
                saved_filename=toc_list_filename,
                file_path=toc_list_path,
                file_size=toc_file_size,
                file_type='toc_list',
                reference_folder=f'scheme_{scheme_id}'
            )
            
            # 保存文档文件引用
            DocumentDAO.save_document_reference(
                task_id=task_id,
                scheme_id=int(scheme_id),
                original_filename=document_filename,
                saved_filename=document_filename,
                file_path=document_path,
                file_size=doc_file_size,
                file_type='document',
                reference_folder=f'scheme_{scheme_id}'
            )
            
            logger.info(f"任务和文档信息已保存到数据库，任务ID: {task_id}")
            
        except Exception as db_error:
            logger.error(f"保存任务到数据库失败: {str(db_error)}")
            return jsonify({
                'code': 500,
                'message': f'保存任务失败: {str(db_error)}',
                'data': {'result': 'false'}
            }), 500
        
        # 准备异步任务参数
        task_params = {
            'task_id': task_id,
            'scheme_id': scheme_id,
            'scheme_name': scheme_name,
            'callback_url': callback_url_full,
            'toc_list_path': toc_list_path,
            'document_path': document_path,
            'toc_list_filename': toc_list_filename,
            'document_filename': document_filename,
            'check_mode': check_mode,
            'embedding_model': embedding_model,
            'chat_model': chat_model,
            'top_k': top_k,
            'openai_api_key': openai_api_key,
            'openai_api_base': openai_api_base,
            'timestamp': timestamp
        }
        
        # 启动异步任务
        thread = threading.Thread(
            target=async_structure_check_worker,
            args=(task_params,),
            daemon=True
        )
        thread.start()
        
        logger.info(f"异步结构检查任务已启动，任务ID: {task_id}")
        
        # 立即返回成功响应
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': {
                'result': 'true',
                'task_id': task_id,
                'callback_url': callback_url_full,
                'estimated_time': '预计3-10分钟内完成'
            }
        }), 200
        
    except Exception as e:
        logger.error(f"异步结构检查API发生错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器内部错误: {str(e)}',
            'data': {'result': 'false'}
        }), 500

# 任务状态查询API
@api_async_structure_check.route('/async_structure_check/status/<task_id>', methods=['GET'])
@swag_from(get_task_status_swagger)
def get_task_status(task_id):
    """
    查询异步任务状态
    ---
    tags:
      - 异步结构检查
    parameters:
      - name: task_id
        in: path
        type: string
        required: true
        description: 任务ID
    responses:
      200:
        description: 任务状态信息
        schema:
          type: object
          properties:
            code:
              type: integer
              example: 200
            message:
              type: string
              example: success
            data:
              type: object
              properties:
                task_id:
                  type: string
                status:
                  type: string
                created_time:
                  type: string
                updated_time:
                  type: string
                result_data:
                  type: object
      404:
        description: 任务不存在
    """
    try:
        task = AsyncTaskDAO.get_task_by_id(task_id)
        if not task:
            return jsonify({
                'code': 404,
                'message': '任务不存在',
                'data': None
            }), 404
        
        task_data = task.to_dict()
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': task_data
        }), 200
        
    except Exception as e:
        logger.error(f"查询任务状态失败: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'查询失败: {str(e)}',
            'data': None
        }), 500

# 检查结果查询API
@api_async_structure_check.route('/async_structure_check/result/<task_id>', methods=['GET'])
@swag_from(get_check_result_swagger)
def get_check_result(task_id):
    """
    获取结构检查结果
    ---
    tags:
      - 异步结构检查
    parameters:
      - name: task_id
        in: path
        type: string
        required: true
        description: 任务ID
    responses:
      200:
        description: 检查结果详情
        schema:
          type: object
          properties:
            code:
              type: integer
              example: 200
            message:
              type: string
              example: success
            data:
              type: object
              properties:
                summary:
                  type: object
                check_results:
                  type: array
                  items:
                    type: object
      404:
        description: 结果不存在
    """
    try:
        result = StructureCheckDAO.get_check_result(task_id)
        if not result:
            return jsonify({
                'code': 404,
                'message': '检查结果不存在',
                'data': None
            }), 404
        
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': result
        }), 200
        
    except Exception as e:
        logger.error(f"获取检查结果失败: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'获取结果失败: {str(e)}',
            'data': None
        }), 500

# 任务列表查询API
@api_async_structure_check.route('/async_structure_check/tasks', methods=['GET'])
@swag_from(list_tasks_swagger)
def list_tasks():
    """
    获取任务列表
    ---
    tags:
      - 异步结构检查
    parameters:
      - name: status
        in: query
        type: string
        description: 任务状态过滤
      - name: limit
        in: query
        type: integer
        default: 20
        description: 返回记录数限制
    responses:
      200:
        description: 任务列表
        schema:
          type: object
          properties:
            code:
              type: integer
              example: 200
            message:
              type: string
              example: success
            data:
              type: object
              properties:
                tasks:
                  type: array
                  items:
                    type: object
                total:
                  type: integer
    """
    try:
        from flask import request
        status = request.args.get('status')
        limit = request.args.get('limit', 20, type=int)
        
        if status:
            tasks = AsyncTaskDAO.get_tasks_by_status(status, hours=24*7, limit=limit)
        else:
            from db.models import AsyncTask
            tasks = AsyncTask.find_all(
                "task_type = 'structure_check' ORDER BY created_time DESC",
                limit=limit
            )
        
        task_list = [task.to_dict() for task in tasks]
        
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': {
                'tasks': task_list,
                'total': len(task_list)
            }
        }), 200
        
    except Exception as e:
        logger.error(f"获取任务列表失败: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'获取任务列表失败: {str(e)}',
            'data': None
        }), 500

# 导入现有的结构检查辅助函数
def perform_item_by_item_structure_check(toc_items, auditor, chat_model, top_k):
    """逐条检查模式"""
    results = []
    
    for i, item in enumerate(toc_items):
        try:
            chapter = item.get('章节', '')
            name = item.get('名称', '')
            required = item.get('必有', '否')
            item_type = item.get('类型', '')
            ai_applicable = item.get('AI适用', '否')
            description = item.get('说明', '')
            
            # 构建检索查询
            search_query = f"{chapter} {name}"
            
            # 检索相关文档片段
            try:
                similar_chunks_results = auditor.search_similar_chunks(search_query, top_k=top_k)
                
                # 过滤相似度低的结果
                similar_chunks = []
                for result in similar_chunks_results:
                    if result['similarity'] > 0.1:
                        similar_chunks.append((result['text'], result['similarity']))
                
                # 构建证据文本
                evidence_text = ""
                if similar_chunks:
                    evidence_chunks = []
                    for chunk, similarity in similar_chunks:
                        evidence_chunks.append(f"相关度{similarity:.3f}: {chunk}")
                    evidence_text = '\n'.join(evidence_chunks[:3])  # 最多3个相关片段
                
                # 使用AI分析结构完整性
                if ai_applicable == '是' and evidence_text:
                    analysis_result = analyze_structure_completeness_single(
                        item, evidence_text, chat_model, auditor
                    )
                    analysis_result['item_id'] = str(i + 1)
                    results.append(analysis_result)
                else:
                    # 对于不适用AI的项目，进行简单的关键词匹配
                    simple_result = simple_structure_check_single(item, evidence_text, i + 1)
                    results.append(simple_result)
                    
            except Exception as e:
                logger.error(f"检索文档片段失败: {str(e)}")
                results.append({
                    'item_id': str(i + 1),
                    'chapter': chapter,
                    'name': name,
                    'required': required,
                    'item_type': item_type,
                    'ai_applicable': ai_applicable,
                    'description': description,
                    'completeness_status': '检查失败',
                    'completeness_score': 0.0,
                    'evidence': '',
                    'detailed_result': f'检索失败: {str(e)}'
                })
                
        except Exception as e:
            logger.error(f"检查项目 {i+1} 时发生错误: {str(e)}")
            results.append({
                'item_id': str(i + 1),
                'chapter': item.get('章节', ''),
                'name': item.get('名称', ''),
                'required': item.get('必有', '否'),
                'item_type': item.get('类型', ''),
                'ai_applicable': item.get('AI适用', '否'),
                'description': item.get('说明', ''),
                'completeness_status': '检查失败',
                'completeness_score': 0.0,
                'evidence': '',
                'detailed_result': f'检查过程出错: {str(e)}'
            })
    
    return results

def perform_chapter_by_chapter_structure_check(toc_items, auditor, chat_model, top_k):
    """逐章节检查模式"""
    results = []
    
    # 按章节分组
    chapters = {}
    for i, item in enumerate(toc_items):
        chapter_num = item.get('章节', '')
        # 提取主章节号（如"1.1.2"的"1"）
        chapter_prefix = chapter_num.split('.')[0] if '.' in chapter_num else chapter_num
        if chapter_prefix not in chapters:
            chapters[chapter_prefix] = []
        chapters[chapter_prefix].append((i, item))
    
    # 逐章节检查
    for chapter_prefix, chapter_items in chapters.items():
        try:
            # 构建章节级检索查询
            chapter_query_parts = []
            for _, item in chapter_items:
                chapter_query_parts.append(f"{item.get('章节', '')} {item.get('名称', '')}")
            
            chapter_query = f"第{chapter_prefix}章 " + " ".join(chapter_query_parts[:5])  # 限制查询长度
            
            # 检索相关文档片段
            try:
                logger.info(f"开始检索章节 {chapter_prefix}，查询: {chapter_query}")
                
                # 验证auditor对象和FAISS索引
                if not auditor or not hasattr(auditor, 'search_similar_chunks'):
                    raise Exception(f"auditor对象无效或缺少search_similar_chunks方法")
                
                if not hasattr(auditor, 'faiss_index') or auditor.faiss_index is None:
                    raise Exception(f"FAISS索引未初始化")
                
                logger.debug(f"FAISS索引维度: {auditor.faiss_index.d if hasattr(auditor.faiss_index, 'd') else 'unknown'}")
                
                search_top_k = min(top_k * 3, 20)
                logger.debug(f"调用search_similar_chunks，查询长度: {len(chapter_query)}, top_k={search_top_k}")
                
                try:
                    similar_chunks_results = auditor.search_similar_chunks(chapter_query, top_k=search_top_k)
                except Exception as search_error:
                    if "assert d == self.d" in str(search_error) or "AssertionError" in str(search_error):
                        logger.error(f"向量维度不匹配错误，强制重新初始化auditor")
                        
                        # 强制清理缓存目录
                        import shutil
                        if hasattr(auditor, 'cache_dir') and os.path.exists(auditor.cache_dir):
                            logger.info(f"删除缓存目录: {auditor.cache_dir}")
                            shutil.rmtree(auditor.cache_dir, ignore_errors=True)
                        
                        # 重新初始化auditor对象
                        logger.info("重新初始化PlanAuditor对象")
                        auditor = PlanAuditor(
                            plan_content=auditor.plan_content if hasattr(auditor, 'plan_content') else doc_text,
                            check_list_file=temp_check_list,
                            embedding_model=embedding_model,
                            openai_api_key=openai_api_key,
                            openai_api_base=openai_api_base,
                            cache_dir=scheme_cache_dir,
                            original_filename=document_filename
                        )
                        
                        # 重新构建索引
                        logger.info("重新构建嵌入索引")
                        plan_id = auditor.build_or_load_embeddings(use_cache=False)
                        logger.info(f"索引重新构建完成，plan_id: {plan_id}")
                        
                        # 重试搜索
                        logger.info("重试搜索操作")
                        similar_chunks_results = auditor.search_similar_chunks(chapter_query, top_k=search_top_k)
                        logger.info("重试搜索成功")
                    else:
                        logger.error(f"其他类型的搜索错误: {str(search_error)}")
                        raise search_error
                
                # 验证返回结果
                if similar_chunks_results is None:
                    logger.warning(f"章节 {chapter_prefix} search_similar_chunks返回none")
                    similar_chunks_results = []
                logger.info(f"章节 {chapter_prefix} 检索到 {len(similar_chunks_results) if similar_chunks_results else 0} 个结果")
                
                # 过滤相似度低的结果
                similar_chunks = []
                if similar_chunks_results:
                    for idx, result in enumerate(similar_chunks_results):
                        try:
                            if isinstance(result, dict) and 'similarity' in result and 'text' in result:
                                if result['similarity'] > 0.1:
                                    similar_chunks.append((result['text'], result['similarity']))
                            else:
                                logger.warning(f"章节 {chapter_prefix} 第{idx}个检索结果格式异常: {type(result)} - {result}")
                        except Exception as result_error:
                            logger.error(f"章节 {chapter_prefix} 处理第{idx}个结果时出错: {str(result_error)}")
                
                logger.info(f"章节 {chapter_prefix} 过滤后有效结果: {len(similar_chunks)} 个")
                
                # 构建章节证据文本
                evidence_text = ""
                if similar_chunks:
                    evidence_chunks = []
                    for chunk, similarity in similar_chunks:
                        evidence_chunks.append(f"相关度{similarity:.3f}: {chunk}")
                    evidence_text = '\n'.join(evidence_chunks[:10])  # 章节级分析用更多证据
                
                logger.info(f"章节 {chapter_prefix} 证据文本长度: {len(evidence_text)}")
                
                # 使用AI进行章节级批量分析
                logger.info(f"开始AI分析章节 {chapter_prefix}")
                chapter_analysis = analyze_chapter_structure_completeness_batch(
                    chapter_items, evidence_text, chat_model, auditor
                )
                
                if chapter_analysis is None:
                    logger.error(f"章节 {chapter_prefix} AI分析返回none")
                    raise Exception(f"章节AI分析失败")
                
                logger.info(f"章节 {chapter_prefix} AI分析完成")
                
                # 将章节分析结果分配到各个项目
                for i, item in chapter_items:
                    try:
                        item_result = extract_item_result_from_chapter_analysis(
                            i + 1, item, chapter_analysis, evidence_text
                        )
                        results.append(item_result)
                    except Exception as item_error:
                        logger.error(f"章节 {chapter_prefix} 项目 {i+1} 结果提取失败: {str(item_error)}")
                        # 使用简单检查作为备用
                        fallback_result = simple_structure_check_single(item, evidence_text, i + 1)
                        results.append(fallback_result)
                
                logger.info(f"章节 {chapter_prefix} 处理完成，共 {len(chapter_items)} 个项目")
                    
            except Exception as e:
                import traceback
                logger.error(f"检索章节 {chapter_prefix} 失败: {str(e)}")
                logger.error(f"异常详情: {traceback.format_exc()}")
                # 章节检查失败时，为该章节的所有项目添加失败结果
                for i, item in chapter_items:
                    results.append({
                        'item_id': str(i + 1),
                        'chapter': item.get('章节', ''),
                        'name': item.get('名称', ''),
                        'required': item.get('必有', '否'),
                        'item_type': item.get('类型', ''),
                        'ai_applicable': item.get('AI适用', '否'),
                        'description': item.get('说明', ''),
                        'completeness_status': '检查失败',
                        'completeness_score': 0.0,
                        'evidence': '',
                        'detailed_result': f'章节检索失败: {str(e)}'
                    })
                
        except Exception as e:
            import traceback
            logger.error(f"检查章节 {chapter_prefix} 时发生错误: {str(e)}")
            logger.error(f"异常详情: {traceback.format_exc()}")
            # 章节检查失败时，为该章节的所有项目添加失败结果
            for i, item in chapter_items:
                results.append({
                    'item_id': str(i + 1),
                    'chapter': item.get('章节', ''),
                    'name': item.get('名称', ''),
                    'required': item.get('必有', '否'),
                    'item_type': item.get('类型', ''),
                    'ai_applicable': item.get('AI适用', '否'),
                    'description': item.get('说明', ''),
                    'completeness_status': '检查失败',
                    'completeness_score': 0.0,
                    'evidence': '',
                    'detailed_result': f'章节检查过程出错: {str(e)}'
                })
    
    # 按原始顺序排序
    results.sort(key=lambda x: int(x['item_id']))
    return results

def analyze_structure_completeness_single(item, evidence_text, chat_model, auditor):
    """使用AI分析单个项目的结构完整性"""
    chapter = item.get('章节', '')
    name = item.get('名称', '')
    required = item.get('必有', '否')
    item_type = item.get('类型', '')
    description = item.get('说明', '')
    
    messages = [
        {
            "role": "system",
            "content": "你是一位专业的建筑工程施工方案审查专家。请分析文档中是否包含指定的目录项内容，并评估其完整性。"
        },
        {
            "role": "user",
            "content": f"""请分析文档中是否包含指定的目录项内容。

目录项信息：
- 章节：{chapter}
- 名称：{name}
- 是否必有：{required}
- 类型：{item_type}
- 检查说明：{description}

文档相关内容：
{evidence_text}

请按以下格式回答：
1. 完整性状态：[完整/部分完整/缺失]
2. 完整性评分：[0.0-1.0之间的数值]
3. 分析说明：[详细说明分析过程和判断依据]

评估标准：
- 完整：目录项内容齐全，满足施工方案要求
- 部分完整：目录项内容存在但不够完整或详细
- 缺失：未找到相关内容或内容严重不足
"""
        }
    ]
    
    try:
        llm_response = auditor.embedder.generate_text(
            messages=messages,
            model=chat_model,
            temperature=0.1
        )
        
        # 解析AI返回结果
        completeness_status = "缺失"  # 默认状态
        completeness_score = 0.0  # 默认评分
        
        # 简单的文本解析来提取完整性状态
        response_lower = llm_response.lower()
        if "完整" in response_lower and "部分" not in response_lower:
            completeness_status = "完整"
        elif "部分完整" in response_lower or "部分" in response_lower:
            completeness_status = "部分完整"
        elif "缺失" in response_lower:
            completeness_status = "缺失"
        
        # 尝试提取完整性评分
        import re
        patterns = [
            r'完整性评分[：:]\s*([0-9.]+)',
            r'评分[：:]\s*([0-9.]+)',
            r'2\.\s*完整性评分[：:]\s*([0-9.]+)'
        ]
        
        for pattern in patterns:
            score_match = re.search(pattern, llm_response)
            if score_match:
                try:
                    completeness_score = float(score_match.group(1))
                    completeness_score = max(0.0, min(1.0, completeness_score))  # 确保在0-1范围内
                    break
                except ValueError:
                    continue
        
        return {
            'chapter': chapter,
            'name': name,
            'required': required,
            'item_type': item_type,
            'ai_applicable': item.get('AI适用', '否'),
            'description': description,
            'completeness_status': completeness_status,
            'completeness_score': round(completeness_score, 3),
            'evidence': evidence_text[:500] + '...' if len(evidence_text) > 500 else evidence_text,
            'detailed_result': llm_response
        }
        
    except Exception as e:
        logger.error(f"AI分析过程出错: {str(e)}")
        return simple_structure_check_single(item, evidence_text, 0)

def simple_structure_check_single(item, evidence_text, item_id):
    """简单的结构完整性检查（关键词匹配）"""
    chapter = item.get('章节', '')
    name = item.get('名称', '')
    required = item.get('必有', '否')
    
    # 基于关键词匹配的简单检查
    keywords = [name]
    if chapter:
        keywords.append(chapter)
    
    found_keywords = []
    for keyword in keywords:
        if keyword and keyword in evidence_text:
            found_keywords.append(keyword)
    
    if found_keywords:
        if len(found_keywords) == len([k for k in keywords if k]):
            status = '完整'
            score = 0.8
        else:
            status = '部分完整'
            score = 0.5
    else:
        if required == '是':
            status = '缺失'
            score = 0.0
        else:
            status = '部分完整'  # 非必须项目给予部分完整
            score = 0.3
    
    return {
        'item_id': str(item_id),
        'chapter': chapter,
        'name': name,
        'required': required,
        'item_type': item.get('类型', ''),
        'ai_applicable': item.get('AI适用', '否'),
        'description': item.get('说明', ''),
        'completeness_status': status,
        'completeness_score': score,
        'evidence': evidence_text[:500] + '...' if len(evidence_text) > 500 else evidence_text,
        'detailed_result': f"基于关键词匹配的简单检查。找到关键词: {', '.join(found_keywords) if found_keywords else '无'}"
    }

def analyze_chapter_structure_completeness_batch(chapter_items, evidence_text, chat_model, auditor):
    """使用AI分析整个章节的结构完整性"""
    try:
        logger.info(f"开始AI批量分析，项目数: {len(chapter_items)}, 证据长度: {len(evidence_text)}")
        
        chapter_info = []
        for i, item in chapter_items:
            chapter_info.append({
                'item_id': i + 1,
                'chapter': item.get('章节', ''),
                'name': item.get('名称', ''),
                'required': item.get('必有', '否'),
                'type': item.get('类型', ''),
                'description': item.get('说明', '')
            })
        
        logger.info(f"章节信息数据构建完成: {len(chapter_info)} 个项目")
        
        messages = [
        {
            "role": "system",
            "content": "你是一位专业的建筑工程施工方案审查专家。请分析文档中指定章节的结构完整性，逐项评估每个目录项。"
        },
        {
            "role": "user",
            "content": f"""请分析文档中指定章节的结构完整性。

需要检查的章节项目：
{json.dumps(chapter_info, ensure_ascii=False, indent=2)}

文档相关内容：
{evidence_text}

请逐项分析每个目录项的完整性状况，并按以下格式回答：

章节整体分析：[对整个章节的总体评价]

逐项分析：
项目1 (ID: X):
- 完整性状态：[完整/部分完整/缺失]
- 完整性评分：[0.0-1.0之间的数值]
- 分析说明：[具体分析]

项目2 (ID: Y):
...

评估标准：
- 完整：目录项内容齐全，满足施工方案要求
- 部分完整：目录项内容存在但不够完整或详细
- 缺失：未找到相关内容或内容严重不足
"""
        }
        ]
        
        try:
            logger.info(f"开始调用LLM进行AI分析")
            llm_response = auditor.embedder.generate_text(
                messages=messages,
                model=chat_model,
                temperature=0.1
            )
            
            logger.info(f"LLM分析完成，响应长度: {len(llm_response)}")
            
            return {
                'chapter_analysis': llm_response,
                'raw_response': llm_response
            }
            
        except Exception as e:
            import traceback
            logger.error(f"章节AI分析过程出错: {str(e)}")
            logger.error(f"异常详情: {traceback.format_exc()}")
            return None
    
    except Exception as e:
        import traceback
        logger.error(f"整个章节分析函数出错: {str(e)}")
        logger.error(f"异常详情: {traceback.format_exc()}")
        return None

def extract_item_result_from_chapter_analysis(item_id, item, chapter_analysis, evidence_text):
    """从章节分析结果中提取单个项目的结果"""
    try:
        logger.debug(f"提取项目 {item_id} 的分析结果")
        
        if not chapter_analysis:
            logger.warning(f"项目 {item_id} 没有章节分析结果")
            return simple_structure_check_single(item, evidence_text, item_id)
        
        if chapter_analysis and 'raw_response' in chapter_analysis:
            response = chapter_analysis['raw_response']
        
        # 尝试解析逐项分析结果
        import re
        
        # 查找对应项目的分析
        pattern = rf'项目.*?\(?ID:?\s*{item_id}\)?:?\s*(.*?)(?=项目.*?\(?ID:?\s*\d+\)?:|$)'
        match = re.search(pattern, response, re.DOTALL | re.IGNORECASE)
        
        if match:
            item_analysis = match.group(1).strip()
            
            # 解析完整性状态
            status = "缺失"
            if "完整性状态" in item_analysis:
                if "完整" in item_analysis and "部分" not in item_analysis:
                    status = "完整"
                elif "部分完整" in item_analysis or "部分" in item_analysis:
                    status = "部分完整"
                elif "缺失" in item_analysis:
                    status = "缺失"
            
            # 解析评分
            score = 0.0
            score_patterns = [
                r'完整性评分[：:]\s*([0-9.]+)',
                r'评分[：:]\s*([0-9.]+)'
            ]
            
            for pattern in score_patterns:
                score_match = re.search(pattern, item_analysis)
                if score_match:
                    try:
                        score = float(score_match.group(1))
                        score = max(0.0, min(1.0, score))  # 确保在0-1范围内
                        break
                    except ValueError:
                        continue
            
            return {
                'item_id': str(item_id),
                'chapter': item.get('章节', ''),
                'name': item.get('名称', ''),
                'required': item.get('必有', '否'),
                'item_type': item.get('类型', ''),
                'ai_applicable': item.get('AI适用', '否'),
                'description': item.get('说明', ''),
                'completeness_status': status,
                'completeness_score': round(score, 3),
                'evidence': evidence_text[:500] + '...' if len(evidence_text) > 500 else evidence_text,
                'detailed_result': item_analysis
            }
    
        # 如果没有找到对应的分析结果，使用简单检查
        logger.debug(f"项目 {item_id} 未找到对应分析结果，使用简单检查")
        return simple_structure_check_single(item, evidence_text, item_id)
    
    except Exception as e:
        import traceback
        logger.error(f"提取项目 {item_id} 结果时出错: {str(e)}")
        logger.error(f"异常详情: {traceback.format_exc()}")
        return simple_structure_check_single(item, evidence_text, item_id)

# 本地测试回调接口
@api_async_structure_check.route('/test/callback', methods=['POST'])
def test_callback():
    """
    本地测试回调接口
    ---
    tags:
      - 测试接口
    summary: 接收异步任务回调的测试接口
    description: 用于本地测试的回调接口，接收结构检查任务的回调数据
    parameters:
      - name: body
        in: body
        required: true
        description: 回调数据
        schema:
          type: object
          properties:
            task_id:
              type: string
              description: 任务ID
            status:
              type: string
              description: 任务状态
            timestamp:
              type: string
              description: 时间戳
            data:
              type: object
              description: 检查结果数据
            error_message:
              type: string
              description: 错误信息
    responses:
      200:
        description: 回调接收成功
        schema:
          type: object
          properties:
            code:
              type: integer
              example: 200
            message:
              type: string
              example: success
            data:
              type: object
      400:
        description: 请求参数错误
    """
    try:
        # 获取回调数据
        callback_data = request.get_json()
        
        if not callback_data:
            logger.warning("回调接收失败：未收到有效的JSON数据")
            return jsonify({
                'code': 400,
                'message': '无效的回调数据',
                'data': None
            }), 400
        
        task_id = callback_data.get('task_id')
        status = callback_data.get('status')
        timestamp = callback_data.get('timestamp')
        data = callback_data.get('data')
        error_message = callback_data.get('error_message')
        
        logger.info(f"收到回调请求 - 任务ID: {task_id}, 状态: {status}, 时间: {timestamp}")
        
        # 打印回调数据详情
        if status == "success" and data:
            summary = data.get('summary', {})
            logger.info(f"任务 {task_id} 执行成功:")
            logger.info(f"  - 总项目数: {summary.get('total_items', 0)}")
            logger.info(f"  - 完整项目数: {summary.get('complete_items', 0)}")
            logger.info(f"  - 缺失项目数: {summary.get('missing_items', 0)}")
            logger.info(f"  - 部分完整项目数: {summary.get('partial_items', 0)}")
            logger.info(f"  - 完整性比例: {summary.get('completeness_rate', 0)}%")
        elif status == "failed":
            logger.error(f"任务 {task_id} 执行失败: {error_message}")
        elif status == "processing":
            logger.info(f"任务 {task_id} 正在处理中...")
        
        # 模拟处理回调数据
        response_data = {
            'received_task_id': task_id,
            'received_status': status,
            'received_at': datetime.now().isoformat(),
            'callback_processed': True
        }
        
        return jsonify({
            'code': 200,
            'message': 'success',
            'data': response_data
        }), 200
        
    except Exception as e:
        logger.error(f"处理回调请求时发生错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'处理回调失败: {str(e)}',
            'data': None
        }), 500
